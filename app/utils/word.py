""" service for creator docx files """
import traceback
from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches



def sjh_full_ver(document_detail: dict, acc, sign):
    doc = Document()
    print(sign)
    
    cover_docx(doc, document_detail['detail_umkm'], acc, sign)
    doc.add_page_break()
    
    base_docx(doc, document_detail['detail_umkm'], document_detail['penetapan_tim'], acc, sign)
    doc.add_page_break()
    
    materi_pelatihan_docx(doc)
    doc.add_page_break()
    
    bukti_pelaksanan_docx(doc, document_detail['bukti_pelaksanaan'], sign)
    doc.add_page_break()
    
    soal_evaluasi_docx(doc, document_detail['jawaban_evaluasi'])
    doc.add_page_break()
    
    audit_internal_docx(doc, document_detail['jawaban_audit'], sign)
    doc.add_page_break()
    
    data_hadir_kaji_ulang_docx(doc, document_detail['daftar_hasil_kaji'], sign)
    doc.add_page_break()
    
    surat_pernyataan_daftar_alamat_docx(doc, acc, sign)
    doc.add_page_break()
    
    form_pembelian_pemeriksaan_bahan_docx(doc, document_detail['pembelian'], document_detail['pembelian_import'])
    doc.add_page_break()
    
    form_stok_bahan_docx(doc, document_detail['stok_barang'])
    doc.add_page_break()
    
    form_produksi_docx(doc, document_detail['form_produksi'])
    doc.add_page_break()
    
    form_pemusnahan_barang_docx(doc, document_detail['form_pemusnahan'])
    doc.add_page_break()
    
    form_pengecheckan_kebersihan_docx(doc, document_detail['form_pengecekan_kebersihan'])
    doc.add_page_break()
    
    daftar_bahan_halal_docx(doc, document_detail['daftar_bahan_halal'], acc, sign)
    doc.add_page_break()
    
    matrik_produk_docx(doc, document_detail['matriks_produk'], acc, sign)
    
    doc.save('./app/assets/'+document_detail['creator']+'.docx')
    return './app/assets/'+document_detail['creator']+'.docx'
    
def cover_docx(doc: Document, detail_umkm, company, sign_data):
    add_head(doc, "MANUAL SJH", 36)
    add_head(doc, "SISTEM JAMINAN HALAL", 28, False)
    add_head(doc, company['company_name'], 28, False)
    
    add_linespace(doc)
    add_linespace(doc)
    add_linespace(doc)
    doc.add_picture(f'./app/assets/{detail_umkm["logo_perusahaan"]}', width=Inches(3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_linespace(doc)
    add_linespace(doc)
    
    doc.add_paragraph(style=doc.styles['Body Text']).add_run('Dibuat dan disahkan oleh :')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_linespace(doc)
    doc.add_picture(f'./app/assets/{sign_data["sign"]}', width=Inches(0.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_head(doc, detail_umkm['nama_ketua'])
    doc.add_paragraph(style=doc.styles['Body Text']).add_run('Tanggal Pengesahan :')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(style=doc.styles['Body Text']).add_run(date.today().strftime("%d %B %Y"))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def base_docx(doc: Document, detail_umkm, penetapan_tim, company, sign_data):
    pedoman_head = "PEDOMAN PENGISIAN TEMPLATE MANUAL SJH"
    pedoman_body = """
    Template Manual SJH merupakan format Manual SJH yang dibuat LPPOM MUI khusus untuk perusahaan kecil. Perusahaan diharapkan dapat memahami isi template Manual SJH agar dapat diterapkan di perusahaan. Perhatikan huruf yang berwarna merah agar disesuaikan dengan kondisi di perusahaan.
    Berikut adalah hal yang perlu diperhatikan dalam penyusunan Manual SJH:"""
    pedoman_list = [{
        "parent": "Informasi Umum Perusahaan pada Pendahuluan",
        "child": [
            "a.	Isi data sesuai dengan perusahaan yang mengajukan sertifikasi halal.", 
            "b.	Jika lokasi pabrik atau fasilitas (dapur/outlet restoran) berbeda dengan lokasi perusahaan, dapat ditambahkan sesuai dengan alamat pabrik dan dapur/outlet."
            ]
    },
    {
        "parent": "Kriteria Sistem Jaminan Halal:",
        "child": [
            "a.	Kebijakan Halal: Isi “….........., ..........................” dengan tempat dan tanggal/bulan/tahun sesuai dengan waktu penetapan Kebijakan Halal. Contoh : Bogor, 5 Januari 2014",
            "b.	Tim Manajemen Halal: disesuaikan dengan kondisi perusahaan. 1 orang tim halal bisa merangkap beberapa posisi.",
            "c.	Fasilitas produksi: disesuaikan dengan apa saja fasilitas yang digunakan pada produksi halal.",
            "d.	Prosedur aktivitas kritis: disesuaikan dengan kondisi perusahaan. Persetujuan bahan baru di LPPOM MUI Provinsi disesuaikan dengan provinsi dimana dilakukan pendaftaran sertifikasi halal.",
            "e.	Laporan berkala audit internal di LPPOM MUI Provinsi disesuaikan dengan provinsi dimana dilakukan pendaftaran sertifikasi halal."
        ]
    },
    {
        "parent": "Lampiran Manual SJH:",
        "child": [
            "a.	Ada 2 jenis lampiran: (i) Lampiran default dari template Manual SJH sebagai bahan rujukan tetapi dapat dimodifikasi oleh perusahaan; (ii) Lampiran tambahan jika ada dokumen yang ingin dilampirkan oleh perusahaan, misalnya Daftar Bahan, jadwal pelatihan, jadwal audit internal, jadwal kaji ulang manajemen, daftar bahan baru.",
            "b.	Lampiran 1: perusahaan dapat memodifikasi isi poster sesuai dengan kebutuhan perusahaan.",
            "c.	Lampiran 2: isi nama lengkap, jabatan di perusahaan serta ditandatangani. Jabatan dapat dirangkap, misalnya 1 orang merangkap sebagai Ketua Tim, R&D dan Pembelian.",
            "d.	Lampiran 3: tidak ada yang perlu diisi. Perusahaan dapat melengkapi materi pelatihan dengan pengetahuan SJH lain yang dianggap perlu untuk diketahui oleh karyawan.",
            "e.	Lampiran 4: isi data bahan yang digunakan untuk seluruh produk yang disertifikasi.",
            "f.	Lampiran 5: isi sesuai formula/resep produk.",
            "g.	Lampiran 6: isi form pemeriksaan bahan setelah bahan dibeli atau bahan diterima dari supplier. Pastikan bahan sesuai dengan yang terdapat di Daftar Bahan Halal. Jika tidak sesuai, maka lakukan evaluasi bahan baru mengikuti prosedur Pemilihan Bahan Baru.",
            "h.	Lampiran 7: daftar pertanyaan pada saat audit internal. Perusahaan dapat memodifikasi daftar pertanyaan sesuai dengan kebutuhan perusahaan.",
            "i.	Lampiran 8: isi form sesuai dengan pembahasan pada rapat kaji ulang manajemen."
        ]
    }
    ]
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.26)
        section.top_margin = Inches(1.06)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.8)

    section = doc.sections[0]

    # ? ============================================================
    # ! Header 
    add_head(doc, pedoman_head, 11)
    body = doc.add_paragraph(style=doc.styles['Body Text']).add_run(f'{pedoman_body}')
    body.font.size = Pt(10)


    for _list in pedoman_list:
        body = doc.add_paragraph(style='List Number').add_run(f'{_list["parent"]}')
        body.font.size = Pt(10)
        for grand_child in _list["child"]:
            body = doc.add_paragraph(style='List 2').add_run(f'{grand_child}')
            body.font.size = Pt(10)


    # * ============================================================
    doc.add_page_break()
    # * ============================================================

    pendahuluan_head = "PENDAHULUAN"
    pendahuluan_body = [{
        "parent": "I. Informasi Umum Perusahaan",
        "child": [
            f"Nama Perusahaan \t\t: {company['company_name']}",
            f"Alamat Perusahaan \t\t: {company['company_address']}",
            f"Nama Pabrik \t\t: {company['factory_name']}",
            f"Alamat Pabrik  \t\t: {company['factory_address']}",
            f"Telp/Fax Perusahaan \t: {company['company_number']}",
            f"Contact Person/Email \t: {company['email']}",
            f"Nama/Merk Produk \t\t: {company['product_name']}",
            f"Jenis Produk \t\t: {company['product_type']}",
            f"Daerah Pemasaran \t\t: {company['marketing_area']}",
            f"Sistem Pemasaran \t\t: {company['marketing_system']}"
        ]
    },{
        "parent": "II. Tujuan",
        "child": ["Manual SJH disusun untuk menjadi pedoman dalam penerapan SJH di perusahaan, dalam rangka menjaga kesinambungan produksi halal sesuai dengan persyaratan sertifikasi halal MUI."]
    }]
    # ? ============================================================
    # ! Header 
    add_head(doc, pendahuluan_head)

    for data in pendahuluan_body:
        body = doc.add_paragraph(style='List').add_run(f'{data["parent"]}')
        body.font.size = Pt(11)
        body.font.bold = True
        for grand_child in data["child"]:
            body = doc.add_paragraph(style='List 2').add_run(f'{grand_child}')
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body.font.size = Pt(10)
        doc.add_paragraph(style='Body Text').add_run()

    kriteria_head = "KRITERIA SISTEM JAMINAN HALAL"
    kebijakan_halal_head = "I.\tKebijakan Halal"
    kebijakan_halal_card = {
        "header": "KEBIJAKAN HALAL",
        "company_name": company['company_name'],
        "body": '“Kami berkomitmen tinggi untuk menghasilkan produk halal, dengan hanya menggunakan bahan yang telah disetujui oleh LPPOM MUI dan diproduksi dengan menggunakan peralatan yang bebas dari najis. Kami akan mencapainya dengan membentuk tim manajemen halal dan melaksanakan dengan sungguh-sungguh Sistem Jaminan Halal”',
        "date": date.today().strftime("%d %B %Y")
    }
    kebijakan_halal_body = "Kebijakan halal disosialisasikan ke seluruh karyawan dengan menempel Kebijakan halal dan poster halal (Lampiran 1) di lokasi yang mudah dilihat seperti di area kantor dan produksi, serta melalui email ke supplier bahan baku (jika ada supplier)."

    # ? ============================================================
    # ! Header 
    add_head(doc, kriteria_head)

    add_linespace(doc)

    body = doc.add_paragraph(style='List').add_run(f'{kebijakan_halal_head}')
    body.font.size = Pt(11)
    body.font.bold = True

    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{kebijakan_halal_card["header"]}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.font.size = Pt(11)
    head.font.bold = True 

    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{kebijakan_halal_card["company_name"]}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.font.size = Pt(11)
    head.font.bold = True 


    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{kebijakan_halal_card["body"]}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    head.font.size = Pt(11)

    add_linespace(doc)

    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{kebijakan_halal_card["date"]}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.font.size = Pt(11)

    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{sign_data["title"]}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.font.size = Pt(11)

    # signature image & text
    ceo_sign = doc.styles['Body Text']
    doc.add_picture(f'./app/assets/{sign_data["sign"]}', width=Inches(0.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(style=ceo_sign).add_run(f'{sign_data["name"]}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    ceo_sign.font.size = Pt(11)

    add_linespace(doc)

    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{kebijakan_halal_body}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    head.font.size = Pt(11)


    # * ============================================================
    doc.add_page_break()
    # * ============================================================
    tim_manajemen_head = "II.\tTim Manajemen Halal "
    tim_manajemen_body = "Untuk menerapkan SJH dan dalam rangka menjaga konsistensi kehalalan produk, dengan ini ditunjuk Tim Manajemen Halal yang terdiri dari pimpinan perusahaan dan karyawan seperti tercantum dalam Lampiran 2.  Tugas Tim Manajemen Halal sebagai berikut :"

    tim_manajemen_table = (
        ("Ketua Tim/Pimpinan", (
            "1. Bertanggungjawab dalam sosialisasi kebijakan halal",
            "2. Bertanggungjawab dalam penunjukkan tim manajemen halal",
            "3. Bertanggungjawab dalam pelatihan eksternal/internal",
            "4. Bertanggungjawab dalam audit internal dan pengiriman laporan berkala ke LPPOM MUI",
            "5. Mengkoordinasikan kaji ulang manajemen"
            )
        ),
        ("Pembelian", (
            "1. Bertanggungjawab dalam proses seleksi bahan baru",
            "2. Bertanggungjawab dalam proses pembelian bahan baku",
            "3. Membuat dan memperbaharui Daftar Bahan Halal serta memonitoring masa berlaku dokumen bahan",
            "4. Bertanggungjawab dalam proses pemeriksaan kedatangan bahan",
            ),
        ),
        ("Produksi", (
            "1. Bertanggungjawab dalam proses produksi halal",
            "2. Bertanggungjawab dalam proses pengembangan produk",
            "3. Bertanggungjawab dalam proses pencucian fasilitas produksi",
            "4. Bertanggungjawab dalam pengananan produk yang tidak memenuhi kriteria (jika terjadi)",
            ),
        ),
        ("Penyimpanan dan Pengiriman", (
            "1. Bertanggungjawab dalam proses penyimpanan bahan dan produk jadi",
            "2. Bertanggungjawab dalam proses transportasi bahan dan produk jadi"
        ))
    )
    tim_manajemen_note = "*) Jabatan dapat dirangkap, misalnya 1 orang merangkap sebagai Ketua Tim, R&D dan Pembelian"

    body = doc.add_paragraph(style='List').add_run(f'{tim_manajemen_head}')
    body.font.size = Pt(11)
    body.font.bold = True

    body = doc.add_paragraph(style='List').add_run(f'\t{tim_manajemen_body}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.font.size = Pt(11)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Bagian/Jabatan'
    hdr_cells[1].text = 'Tugas Tim Manajemen Halal'
    for bagian, tugas_tugas in tim_manajemen_table:
        row_cells = table.add_row().cells
        row_cells[0].text = bagian
        for tugas in tugas_tugas:
            row_cells[1].text += tugas + "\n"
    table.allow_autofit = True
    for cell in table.columns[0].cells:
        cell.width = Inches(1)
    for cell in table.columns[1].cells:
        cell.width = Inches(5.56)

    body = doc.add_paragraph(style='List 2').add_run(f'{tim_manajemen_note}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.font.size = Pt(11)

    add_linespace(doc)

    pelatihan_head = "III.	Pelatihan"
    pelatihan_list = [
        "a.	Tujuan pelatihan yaitu meningkatkan pemahaman Tim Manajemen Halal tentang persyaratan sertifikasi halal.",
        "b.	Pelatihan dilakukan secara internal dan eksternal.",
        "c.	Pelatihan internal dilakukan 1 tahun sekali oleh personel yang sudah mengikuti pelatihan halal. Pelatihan diikuti oleh seluruh karyawan. Setiap karyawan baru harus mendapatkan pelatihan internal sebelum mulai bekerja. Materi dan soal tes pelatihan internal sesuai dengan Lampiran 3 atau dapat dimodifikasi sesuai kebutuhan. Daftar hadir pelatihan dan hasil tes disimpan selama dua tahun sebagai bukti pelatihan internal.",
        "d.	Salah satu tim manajemen halal akan mengikuti pelatihan eksternal ke lembaga pelatihan halal dengan program pelatihan berbasis HAS 23000. Sertifikat pelatihan eksternal disimpan selama dua tahun sebagai bukti pelatihan eksternal. Pelatihan eksternal dapat diikuti setiap 2 tahun sekali atau sesuai kebutuhan perusahaan."
    ]

    body = doc.add_paragraph(style='List').add_run(f'{pelatihan_head}')
    body.font.size = Pt(11)
    body.font.bold = True

    for pelatihan in pelatihan_list:
        body = doc.add_paragraph(style='List 2').add_run(f'{pelatihan}')
        body.font.size = Pt(11)

    add_linespace(doc)

    bahan_head = "IV.	Bahan"
    bahan_list = [
        "a.	Perusahaan hanya menggunakan bahan yang sesuai dengan kriteria SJH dan disetujui oleh LPPOM MUI untuk menghasilkan produk yang disertifikasi.",
        "b.	Perusahaan membuat Daftar Bahan Halal yaitu daftar bahan yang disetujui LPPOM MUI, yang dapat dilihat pada Lampiran 4.",
        "c.	Daftar Bahan Halal digunakan sebagai acuan dalam proses pembelian, pemeriksaan bahan datang dan proses produksi.",
        "d.	Perbaikan Daftar Bahan Halal dilakukan bila terjadi perubahan bahan.",
        "e.	Semua bahan dilengkapi dengan dokumen pendukung, kecuali bahan tidak kritis (positive list) yang tercantum dalam SK LPPOM MUI. Dokumen pendukung dapat berupa Sertifikat halal atau dokumen lain dari produsen bahan yang menjelaskan sumber bahan.",
        "f.	Ketua tim manajemen halal atau wakilnya akan memeriksa masa berlaku Sertifikat halal yang terdapat di Daftar Bahan Halal secara berkala. Jika masa berlaku Sertifikat halal akan berakhir maka dilakukan pemeriksaan masa berlaku Sertifikat halal di website LPPOM MUI (www.halalmui.org) atau meminta Sertifikat halal terbaru ke supplier. Data Sertifikat halal terbaru akan dimasukkan ke dalam Daftar Bahan Halal."
    ]

    body = doc.add_paragraph(style='List').add_run(f'{bahan_head}')
    body.font.size = Pt(11)
    body.font.bold = True

    for bahan in bahan_list:
        body = doc.add_paragraph(style='List 2').add_run(f'{bahan}')
        body.font.size = Pt(11)

    add_linespace(doc)

    produk_head = "V.	Produk"
    produk_list = [
        "a.	Perusahaan hanya akan memproduksi produk halal dengan nama produk yang tidak menggunakan nama yang mengarah pada sesuatu yang diharamkan atau ibadah yang tidak sesuai dengan syariah Islam, bau/rasa produk tidak mengarah kepada produk haram, bentuk produk tidak berupa anjing atau babi, serta kemasan produk tidak vulgar.",
        "b.	Perusahaan akan membuat Matriks Produk untuk semua produk yang disertifikasi halal sesuai dengan formula/resep pada proses produksi, yang dapat dilihat pada Lampiran 5.",
        "c.	Perusahaan akan mendaftarkan seluruh Produk pangan eceran (retail) dengan merk sama yang beredar di Indonesia. Untuk produk pangan bukan eceran (non retail) yang mempunyai merk/brand dan hanya didaftarkan sebagian, maka kami akan mencantumkan logo halal MUI untuk produk yang disertifikasi."
    ]

    body = doc.add_paragraph(style='List').add_run(f'{produk_head}')
    body.font.size = Pt(11)
    body.font.bold = True

    for produk in produk_list:
        body = doc.add_paragraph(style='List 2').add_run(f'{produk}')
        body.font.size = Pt(11)

    add_linespace(doc)

    fasilitas_produksi_head = "VI.	Fasilitas Produksi"
    fasilitas_produksi_list = [
        "a.	Fasilitas produksi harus bersih dan terhindar dari kontaminasi bahan najis.",
        "b.	Fasilitas produksi mencakup seluruh fasilitas yang digunakan untuk proses produksi seperti gudang penyimpanan, ruang pencucian, ruang produksi, dapur/outlet.",
        "c.	Fasilitas produksi yang kami gunakan untuk menghasilkan produk halal tidak digunakan untuk menghasilkan produk lain yang mengandung bahan babi.",
        "d.	Setiap ada tambahan fasilitas produksi yang digunakan untuk menghasilkan produk, baik milik perusahaan sendiri maupun milik pihak lain, maka akan didaftarkan untuk disertifikasi halal dan menjadi ruang lingkup implementasi Sistem Jaminan Halal."
    ]

    body = doc.add_paragraph(style='List').add_run(f'{fasilitas_produksi_head}')
    body.font.size = Pt(11)
    body.font.bold = True

    for fasilitas_produksi in fasilitas_produksi_list:
        body = doc.add_paragraph(style='List 2').add_run(f'{fasilitas_produksi}')
        body.font.size = Pt(11)

    add_linespace(doc)


    prosedur_aktifitas_head = "VII.	Prosedur Aktifitas Kritis"
    prosedur_aktifitas_body = "Berikut adalah prosedur sederhana yang terkait dengan halal :"
    prosedur_aktifitas_table = (
        ("Pemilihan Bahan Baru", (
            "1. Jika produksi produk halal akan menggunakan bahan baru di luar Daftar Bahan Halal (termasuk bahan lama dengan produsen baru), maka kami akan mencari bahan positif list atau bahan bersertifikat halal MUI yang ada di www.halalmui.org.",
            "2. Jika bahan bukan termasuk kategori diatas, maka harus meminta dokumen bahan ke produsen (dapat berupa pernyataan, spesifikasi, diagram alir proses) atau Sertifikat Halal dari lembaga yang diakui MUI, kemudian meminta persetujuan bahan tersebut ke LPPOM MUI.",
            "3. Pengajuan persetujuan ke LPPOM MUI dilakukan melalui Cerol (menu Inquiry Material Approval) atau melalui email (sesuai dengan email LPPOM MUI Provinsi). Bukti persetujuan bahan baru disimpan setidaknya selama dua tahun.",
            "4. Bahan baru dimasukkan ke dalam Daftar Bahan Halal dan disebarkan kepada seluruh tim halal. Bahan baru dapat digunakan setelah mendapatkan persetujuan tertulis dari LPPOM MUI."
        )),
        ("Pembelian", (
            "1.	Melakukan pembelian bahan dengan nama/merk dan produsen sesuai dengan yang tercantum dalam Daftar Bahan Halal.",
            "2.	Bukti pembelian (nota/kuitansi) disimpan setidaknya selama 6 bulan, kecuali untuk bahan yang jarang dibeli maka disimpan bukti pembelian terakhir."
        )),
        ("Pemeriksaan Bahan Datang", (
            "1.	Memeriksa label bahan pada setiap pembelian atau penerimaan bahan untuk memastikan kesesuaian nama bahan, nama produsen dan negara produsen dengan yang tercantum dalam Daftar Bahan Halal. Bahan yang boleh digunakan hanya bahan yang namanya, nama produsen dan negara produsennya sesuai dengan Daftar Bahan Halal.",
            "2.	Pemeriksaan label bahan dituliskan dalam form pemeriksaan bahan seperti pada Lampiran 6.",
            "3.	Pengecualian untuk bahan positif list tidak dilakukan pemeriksaan label bahan."
        )), 
        ("Produksi", (
            "1.	Membuat formula/resep produk yang akan menjadi acuan/rujukan untuk bagian produksi dalam memproduksi produk.",
            "2.	Melakukan produksi sesuai dengan formula dan hanya menggunakan bahan yang tercantum dalam Daftar Bahan Halal. Catatan produksi disimpan setidaknya selama 6 bulan.",
            "3.	Memastikan peralatan produksi dalam keadaan bersih (bebas dari najis) sebelum digunakan."
        )),
        ("Pencucian Fasilitas Produksi", (
            "1.	Peralatan bersih dan kotor dipisahkan",
            "2.	Melakukan proses pencucian dengan bahan pembersih bebas najis",
            "3.	Memastikan proses pencucian menghilangkan kotoran dan najis"
        )),
        ("Penyimpanan & Transportasi", (
            "1.	Menyimpan bahan dan produk di tempat yang bersih dan menjaganya supaya terhindar dari najis.",
            "2.	Memastikan kendaraan yang digunakan untuk mengangkut bahan dan produk halal dalam kondisi baik dan tidak digunakan untuk mengangkut bahan/produk lain yang diragukan kehalalannya."
        )),
        ("Pengembangan Produk Baru", (
            "1.	Setiap ada proses pengembangan produk/menu baru akan didaftarkan untuk disertifikasi ke LPPOM MUI.",
            "2.	Jika sertifikat halal sudah keluar, maka akan dilakukan proses produksi produk/menu baru sesuai dengan formula/resep baku."
        )),
        ("Pengembangan Fasilitas Baru", (
            "1.	Sebelum diajukan untuk mendapatkan sertifikat halal maka pengembangan fasilitas baru dengan terlebih dahulu telah disahkan/telah resmi mendapatkan legalitas untuk semua bentuk perizinan sesuai peraturan daerah/wilayah hukum negara.",
            "2.	Proses pengembangan fasilitas menjamin bahwa pelaksanaan pengembangan fasilitas baru akan didaftarkan sertifikasi halal sebelum dibuka/ digunakan."
        )),
        ("Aturan Karyawan", (
            "1.	Setiap karyawan menjaga kebersihan diri sebelum dan selama bekerja sehingga tidak mengotori produk yang dihasilkan",
            "2.	Setiap karyawan tidak boleh membawa produk tidak halal ke area produksi.",
            "3.	Setiap karyawan tidak boleh membawa/memelihara hewan peliharaan di area produksi.",
            "4.	Setiap karyawan tidak boleh menggunakan peralatan produksi untuk kepentingan lain, misalnya untuk memasak karyawan atau menyimpan produk tidak halal milik karyawan.",
            "5.	Setiap karyawan wajib mencuci tangan sebelum dan sesudah bekerja"
        ))
    )

    body = doc.add_paragraph(style='List').add_run(f'{prosedur_aktifitas_head}')
    body.font.size = Pt(11)
    body.font.bold = True


    body = doc.add_paragraph(style='List 2').add_run(f'{prosedur_aktifitas_body}')
    body.font.size = Pt(11)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Prosedur'
    hdr_cells[1].text = 'Langkah Kerja'
    for prosedur, langkah_kerja in prosedur_aktifitas_table:
        row_cells = table.add_row().cells
        row_cells[0].text = prosedur
        for langkah in langkah_kerja:
            row_cells[1].text += langkah + "\n"
    table.allow_autofit = True
    for cell in table.columns[0].cells:
        cell.width = Inches(1)
    for cell in table.columns[1].cells:
        cell.width = Inches(5.56)

    add_linespace(doc)

    kemampuan_telusur_head = "VIII.	 Kemampuan Telusur"
    kemampuan_telusur_body = "Perusahaan menjamin produk yang disertifikasi tertelusur berasal dari bahan yang sudah disetujui LPPOM MUI (tercantum dalam Daftar Bahan Halal) dan dibuat di fasilitas produksi yang bebas najis. Kemampuan telusur produk dilakukan melalui catatan pemeriksaan bahan datang dan catatan produksi."

    body = doc.add_paragraph(style='List').add_run(f'{kemampuan_telusur_head}')
    body.font.size = Pt(11)
    body.font.bold = True


    body = doc.add_paragraph(style='List 2').add_run(f'{kemampuan_telusur_body}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.font.size = Pt(11)

    add_linespace(doc)

    penanganan_produk_head = "IX.	Penanganan Produk Yang Tidak Memenuhi Kriteria Halal"
    penanganan_produk_list = [
        "a.	Produk yang tidak memenuhi kriteria adalah produk yang sudah disertifikasi halal yang dibuat dari bahan yang tidak disetujui (tidak tercantum dalam Daftar Bahan Halal) dan/atau diproduksi di fasilitas yang tidak bebas najis.",
        "b.	Jika ditemukan produk yang tidak memenuhi kriteria halal, maka produk yang dihasilkan tidak akan dijual ke konsumen. Produk tersebut selanjutnya akan dimusnahkan.",
        "c.	Bila produk yang tidak memenuhi kriteria sudah terlanjur dijual, maka produk tersebut akan ditarik dari pasaran dan selanjutnya akan dimusnahkan.",
        "d.	Bukti pemusnahan produk harus disimpan (jika terjadi)."
    ]

    body = doc.add_paragraph(style='List').add_run(f'{penanganan_produk_head}')
    body.font.size = Pt(11)
    body.font.bold = True

    for penanganan_produk in penanganan_produk_list:
        body = doc.add_paragraph(style='List 2').add_run(f'{penanganan_produk}')
        body.font.size = Pt(11)

    add_linespace(doc)

    audit_internal_head = "X.	Audit Halal Internal"
    audit_internal_list = [
        "a.	Perusahaan akan melakukan audit internal dengan cara memeriksa pelaksanaan SJH dan mengisi form seperti pada Lampiran 7. Pertanyaan audit internal ditujukan kepada Tim halal sesuai dengan pembagian tugas Tim halal pada bagian II.",
        "b.	Audit internal dilakukan oleh Tim Halal dan tidak boleh mengaudit tugasnya sendiri.",
        "c.	Audit internal dilakukan minimal 6 bulan sekali.",
        "d.	Jika dalam audit internal ditemukan kelemahan, yaitu ada pertanyaan yang dijawab “tidak”, maka akan segera dilakukan perbaikan agar kelemahan tersebut tidak terulang. Bukti perbaikan kelemahan harus disimpan setidaknya selama dua tahun.",
        "e.	Bukti pelaksanaan audit internal disimpan setidaknya selama dua tahun.",
        "f.	Hasil audit internal dikirimkan ke LPPOM MUI melalui Cerol pada menu Regular Report atau melalui email (sesuai dengan email LPPOM MUI Provinsi)."
    ]

    body = doc.add_paragraph(style='List').add_run(f'{audit_internal_head}')
    body.font.size = Pt(11)
    body.font.bold = True

    for audit_internal in audit_internal_list:
        body = doc.add_paragraph(style='List 2').add_run(f'{audit_internal}')
        body.font.size = Pt(11)

    add_linespace(doc)

    kaji_ulang_head = "XI.	Kaji Ulang Manajemen"
    kaji_ulang_list = [
        "a.	Perusahaan akan melakukan rapat kaji ulang manajemen minimal setahun sekali.",
        "b.	Rapat kaji ulang manajemen dihadiri oleh pimpinan perusahaan dan Tim Manajemen Halal, membahas hasil dari audit internal atau hal lain terkait SJH seperti adanya perubahan tim halal, penambahan produk, pergantian bahan.",
        "c.	Hasil rapat kaji ulang manajemen dituliskan dalam form hasil rapat kaji ulang manajemen seperti pada Lampiran 8."
    ]

    body = doc.add_paragraph(style='List').add_run(f'{kaji_ulang_head}')
    body.font.size = Pt(11)
    body.font.bold = True

    for kaji_ulang in kaji_ulang_list:
        body = doc.add_paragraph(style='List 2').add_run(f'{kaji_ulang}')
        body.font.size = Pt(11)

    add_linespace(doc)


    # * ============================================================
    doc.add_page_break()
    # * ============================================================

    poster_kebijakan_halal_head = "A.	Poster Kebijakan Halal"
    poster_kebijakan_halal_card = {
        "header": "KEBIJAKAN HALAL",
        "body": '“Kami berkomitmen tinggi untuk menghasilkan produk halal, dengan hanya menggunakan bahan yang telah disetujui oleh LPPOM MUI dan diproduksi dengan menggunakan peralatan yang bebas dari najis. Kami akan mencapainya dengan membentuk tim manajemen halal dan melaksanakan dengan sungguh-sungguh Sistem Jaminan Halal”',
    }

    body = doc.add_paragraph(style='List').add_run(f'{poster_kebijakan_halal_head}')
    body.font.size = Pt(11)
    body.font.bold = True

    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{poster_kebijakan_halal_card["header"]}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.font.size = Pt(11)
    head.font.bold = True 

    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{company["company_name"]}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.font.size = Pt(11)
    head.font.bold = True 


    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{poster_kebijakan_halal_card["body"]}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    head.font.size = Pt(11)

    add_linespace(doc)

    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(date.today().strftime("%d %B %Y"))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.font.size = Pt(11)

    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{sign_data["title"]}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.font.size = Pt(11)

    # signature image & text
    ceo_sign = doc.styles['Body Text']
    doc.add_picture(f'./app/assets/{sign_data["sign"]}', width=Inches(0.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(style=ceo_sign).add_run(f'{sign_data["name"]}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    ceo_sign.font.size = Pt(11)

    add_linespace(doc)

    poster_halal_haram_head = "B.	Poster Halal Haram"
    poster_halal_haram_card = {
        "header" : "PENGERTIAN HALAL HARAM",
        "list": [
            "	Mengkonsumsi makanan dan minuman yang halal adalah wajib hukumnya bagi orang Islam.",
            "	Pengertian halal haram : (i) Halal adalah Boleh. (ii) Haram adalah sesuatu yang dilarang oleh Allah SWT dengan larangan yang tegas.",
            "	Contoh bahan haram : (i) Babi,termasuk seluruh bagian tubuhnya dan produk turunannya (segar atau olahan), (ii) Khamr (minuman beralkohol), (iii) Hewan sembelihan yang tidak disembelih sesuai syariat Islam, (iv) Darah, (v) Bangkai, (vi) Bagian dari tubuh manusia, (vii) Binatang buas, anjing, amfibi."
        ]
    }

    add_linespace(doc)

    body = doc.add_paragraph(style='List').add_run(f'{poster_halal_haram_head}')
    body.font.size = Pt(11)
    body.font.bold = True

    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{poster_halal_haram_card["header"]}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.font.size = Pt(11)
    head.font.bold = True 

    for list_halal_haram in poster_halal_haram_card["list"]:
        body = doc.add_paragraph(style="List 2").add_run(f'{list_halal_haram}')
        body.font.size = Pt(11)


    poster_praktek_penerapan_head = "C.	Poster Praktek Penerapan SJH"
    poster_praktek_penerapan_card = {
        "header": "PRAKTEK PENERAPAN SJH",
        "list": [
            "	Menjaga fasilitas/peralatan produksi dalam keadaan bersih sebelum dan sesudah digunakan.",
            "	Menjaga kebersihan diri sebelum dan selama bekerja sehingga tidak mengotori produk.",
            "	Tidak boleh membawa produk najis/haram di area produksi.",
            "	Tidak boleh membawa/memelihara hewan peliharaan di area produksi.",
            "	Tidak boleh menggunakan peralatan produksi untuk kepentingan lain.",
            "	Menyimpan bahan dan produk di tempat yang bersih dan menjaga agar terhindar dari najis.",
            "	Memastikan kendaraan yang digunakan untuk mengangkut produk halal dalam kondisi baik dan tidak digunakan untuk mengangkut produk lain yang diragukan kehalalannya."
        ]
    }

    add_linespace(doc)

    body = doc.add_paragraph(style='List').add_run(f'{poster_praktek_penerapan_head}')
    body.font.size = Pt(11)
    body.font.bold = True

    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{poster_praktek_penerapan_card["header"]}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.font.size = Pt(11)
    head.font.bold = True 

    for praktek_penerapan in poster_praktek_penerapan_card["list"]:
        body = doc.add_paragraph(style="List 2").add_run(f'{praktek_penerapan}')
        body.font.size = Pt(11)

    add_linespace(doc)


    surat_penetapan_tim_head = "SURAT PENETAPAN TIM MANAJEMEN HALAL"
    surat_penetapan_tim_text_1 = "Untuk menerapkan Sistem Jaminan Halal dan dalam rangka menjaga konsistensi kehalalan produk, dengan ini ditunjuk Tim Manajemen Halal sebagai berikut :"
    surat_penetapan_tim_text_2 = "*) Jabatan dapat dirangkap, misalnya 1 orang (nama cukup ditulis satu kali) merangkap sebagai Ketua Tim, Pembelian, Produksi, Gudang dll ."
    surat_penetapan_tim_text_3 = "Tim Manajemen Halal telah membaca dan memahami Manual SJH serta akan melaksanakan dengan sungguh-sungguh semua prosedur seperti yang tertulis pada Manual SJH."
    surat_penetapan_tim_text_4 = "Demikian surat penetapan ini dibuat untuk dilaksanakan sebagaimana mestinya."


    # * ============================================================
    doc.add_page_break()
    # * ============================================================

    heading_style = doc.styles['Body Text']
    body = doc.add_paragraph(style=heading_style).add_run('Lampiran 2. Surat Penetapan Tim Manajemen Halal')
    body.font.size = Pt(11)
    body.font.bold = True

    # ? ============================================================
    # ! Header 
    add_linespace(doc)

    add_head(doc, surat_penetapan_tim_head)
    add_linespace(doc)


    heading_style = doc.styles['Body Text']
    body = doc.add_paragraph(style=heading_style).add_run(f'{surat_penetapan_tim_text_1}')
    body.font.size = Pt(11)


    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No.'
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells[1].text = 'Nama'
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells[2].text = 'Jabatan*)'
    hdr_cells[2].paragraphs[0].runs[0].font.bold = True
    hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells[3].text = 'Posisi di TIM'
    hdr_cells[3].paragraphs[0].runs[0].font.bold = True
    hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    index = 1
    for tim in penetapan_tim:
        row_cells = table.add_row().cells
        row_cells[0].text = str(index)
        row_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].text = tim['nama']
        row_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        row_cells[2].text = tim['jabatan']
        row_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        row_cells[3].text = tim['position']
        row_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        index += 1
    for cell in table.columns[0].cells:
        cell.width = Inches(0.56)
    for cell in table.columns[1].cells:
        cell.width = Inches(2)
    for cell in table.columns[2].cells:
        cell.width = Inches(2)
    for cell in table.columns[3].cells:
        cell.width = Inches(2)
    for row in table.rows:
        row.height = Pt(32)

    heading_style = doc.styles['Body Text']
    body = doc.add_paragraph(style=heading_style).add_run(f'{surat_penetapan_tim_text_2}')
    body.font.size = Pt(11)

    heading_style = doc.styles['Body Text']
    body = doc.add_paragraph(style=heading_style).add_run(f'{surat_penetapan_tim_text_3}')
    body.font.size = Pt(11)

    heading_style = doc.styles['Body Text']
    body = doc.add_paragraph(style=heading_style).add_run(f'{surat_penetapan_tim_text_4}')
    body.font.size = Pt(11)

    add_linespace(doc)
    add_linespace(doc)

    heading_style = doc.styles['Body Text']
    body = doc.add_paragraph(style=heading_style).add_run({date.today().strftime("%d %B %Y")})
    body.font.size = Pt(11)

    heading_style = doc.styles['Body Text']
    body = doc.add_paragraph(style=heading_style).add_run(sign_data['title'])
    body.font.size = Pt(11)

    ceo_sign = doc.styles['Body Text']
    doc.add_picture(f'./app/assets/{sign_data["sign"]}', width=Inches(0.7))
    doc.add_paragraph(style=ceo_sign).add_run(sign_data['name'])
    ceo_sign.font.size = Pt(11)


def materi_pelatihan_docx(doc: Document):
    primary_style = doc.styles['Body Text']
    primary_style.font.size = Pt(11)
    
    
    header = doc.add_paragraph(style=primary_style).add_run('Lampiran 3.  Materi Pelatihan Internal dan Soal Tes Pelatihan Internal')
    header.font.bold = True
    add_linespace(doc)
    
    title = doc.add_paragraph(style=primary_style).add_run('DAFTAR HADIR KAJI ULANG MANAJEMEN')
    title.font.bold = True
    
    materi = [
        {
            "head": "A.	Pengetahuan Halal Haram",
            "body": [
                "1.	Mengkonsumsi makanan dan minuman yang halal adalah wajib hukumnya bagi orang Islam.",
                "2.	Pengertian halal haram : (i) Halal adalah Boleh. (ii) Haram adalah sesuatu yang dilarang oleh Allah SWT dengan larangan yang tegas.",
                "3.	Contoh bahan haram : (i) Babi,termasuk seluruh bagian tubuhnya dan produk turunannya (segar atau olahan), (ii) Khamr (minuman beralkohol), (iii) Hasil samping khamr yang diperoleh hanya dengan pemisahan secara fisik, (iv) Darah, (v) Bangkai, (vi) Bagian dari tubuh manusia, binatang buas, anjing."
            ]
        },
        {
            "head": "B.	Pengetahuan Benda Najis",
            "body": [
                "1.	Pengertian najis : (i) Menurut bahasa adalah “setiap yang kotor”, (ii) Menurut syara’ adalah kotoran yang wajib dihindari dan dibersihkan oleh setiap muslim ketika terkena olehnya.",
                "2.	Najis ada tiga: (1) Najis mukhaffafah(najis ringan), yaitu air seni bayi laki-laki sebelum usia dua tahun yang hanya mengonsumsi ASI, (2) Najis mughallazhah (najis berat), yaitu najis babi, anjing atau turunan keduanya, dan (3) Najis mutawassithah (najis sedang), yaitu najis kotoran hewan, khamr (minuman keras)",
                "3.	Mutanajjis adalah benda suci yang terkena najis, dapat berupa bahan, produk atau peralatan produksi. Benda mutanajjis dapat menjadi suci kembali setelah dicuci secara syar’i. ",
                "4.	Pencucian benda mutanajjis padat yang terkena najis mutawassithah secara syar’i yaitu dengan mengucurinya dengan air atau mencucinya di dalam air yang banyak (direndam) hingga hilang rasa, bau dan warna dari bahan najisnya.",
                "5.	Pencucian benda mutanajjis padat yang terkena najis mughallazhah secara syar’i yaitu dicuci tujuh kali dengan air dan salah satunya dengan tanah atau bahan lain yang mempunyai kemampuan menghilangkan rasa, bau dan warna."
            ]
        },
        {
            "head": "C.	Pengetahuan Sertifikasi Halal",
            "body": [
                "1.	Sertifikat halal produk di Indonesia dikeluarkan oleh Majelis Ulama Indonesia (MUI) setelah audit dilakukan oleh Lembaga Pengkajian Makanan, Obat dan Kosmetika (LPPOM MUI).",
                "2.	Perusahaan yang telah mendapatkan sertifikat halal dari MUI harus menjaga kehalalan produknya dengan cara menerapkan Sistem Jaminan Halal (SJH).",
                "3.	Logo halal tidak boleh digunakan oleh perusahaan jika tidak memiliki sertifikat halal MUI."
            ]
        },
        {
            "head": "D.	Penerapan Sistem Jaminan Halal (SJH)",
            "body": [
                "1.	Inti dari penerapan SJH adalah membuat kebijakan halal, membentuk tim manajemen halal dan melaksanakan dengan sungguh-sungguh semua prosedur operasional yang tercantum dalam Manual SJH.",
                "2.	Kebijakan halal adalah komitmen perusahaan untuk menghasilkan produk halal, dengan hanya menggunakan bahan yang telah disetujui oleh LPPOM MUI dan diproduksi dengan menggunakan peralatan yang bebas dari najis.",
                "3.	Bahan yang telah disetujui oleh LPPOM MUI tercantum dalam Daftar Bahan Halal.",
                "4.	Membeli bahan dengan nama/merk dan produsen sesuai dengan yang tercantum dalam Daftar Bahan Halal.",
                "5.	Jika akan menggunakan bahan baru di luar Daftar Bahan Halal (termasuk bahan lama dengan produsen baru), akan meminta persetujuan penggunaan bahan tersebut ke LPPOM MUI.Bahan bersertifikat halal MUI yang ada di www.halalmui.org dan bahan tidak kritis yang tercantum dalam Lampiran 5 maka tidak dilakukan persetujuan bahan baru ke LPPOM MUI. Bahan baru selanjutnya dimasukkan ke dalam Daftar Bahan Halal.",
                "6.	Memeriksa label bahan pada setiap pembelian bahan atau penerimaan bahan untuk memastikan kesesuaian nama bahan, nama produsen dan negara produsen dengan yang tercantum dalam Daftar Bahan Halal. Bahan yang boleh digunakan hanya bahan yang namanya, nama produsen dan negara produsennya sesuai dengan Daftar Bahan Halal.",
                "7.	Dalam proses produksi hanya menggunakan bahan dengan nama/merk dan produsen seperti yang tercantum dalam Daftar Bahan Halal.",
                "8.	Menjaga semua fasilitas dan peralatan produksi dalam keadaan bersih (bebas dari najis) sebelum dan sesudah digunakan.",
                "9.	Setiap pekerja menjaga kebersihan diri sebelum dan selama bekerja sehingga tidak mengotori produk yang dihasilkan.",
                "10.	Setiap pekerja tidak boleh membawa produk tidak halal di area produksi.",
                "11.	Setiap pekerja tidak boleh membawa/memelihara hewan peliharaan di area produksi.",
                "12.	Setiap pekerja tidak boleh menggunakan peralatan produksi untuk kepentingan lain, misalnya untuk memasak karyawan atau menyimpan produk tidak halal milik karyawan.",
                "13.	Menyimpan bahan dan produk di tempat yang bersih dan menjagaagar terhindar dari najis.",
                "14.	Memastikan kendaraan yang digunakan untuk mengangkut produk halal dalam kondisi baik dan tidak digunakan untuk mengangkut produk lain yang diragukan kehalalannya.",
                "15.	Mendaftarkan setiap ada produk baru retail (eceran) dengan merk yang sama untuk disertifikasi halal sebelum dipasarkan. ",
                "16.	Memastikan produk halal tidakmenggunakan nama yang mengarah pada sesuatu yang diharamkan atau ibadah yang tidaksesuai dengan syariah Islam, dan tidak memilikikecenderungan bau atau rasa yang mengarah kepada produk haram; misalnya coklat Valentine, mie setan, minuman rasa bir, roti rasa daging babi, pasta bikini.",
                "17.	Mendaftarkan setiap ada penambahan fasilitas produksi baru untuk disertifikasi halal."
            ]
        }
    ]
    for sub in materi:
        body = doc.add_paragraph(style='List').add_run(f'{sub["head"]}')
        body.font.size = Pt(11)
        body.font.bold = True

        for produk in sub['body']:
            body = doc.add_paragraph(style='List 2').add_run(f'{produk}')
            body.font.size = Pt(11)
        add_linespace(doc)


def soal_evaluasi_docx(doc: Document, data_evaluasi: dict) -> str:
    """ function for generating soal & jawaban evaluasi files """
    data = [
        {
            "id": 1,
            "soal": "1.	Allah SWT memerintahkan Manusia untuk konsumsi makanan yang….",
            "jawaban": {
                "a": "a.	Halal",
                "b": "b.	Thoyib",
                "c": "c.	Kotor",
                "d": "d.	a dan b"
            }
        }, {
            "id": 2,
            "soal": "2.	Berikut makanan dan minuman yang halal adalah",
            "jawaban": {
                "a": "a.	Klepon",
                "b": "b.	Anjing",
                "c": "c.	Babi",
                "d": "d.	Bangkai Ayam"
            }
        }, {
            "id": 3,
            "soal": "3.	Daging Babi dan turunannya merupakan najis",
            "jawaban": {
                "a": "a.	Ringan (mukhaffafah)",
                "b": "b.	Berat (mughallazhah)",
                "c": "c.	Sedang (mutawassithah)",
                "d": "d.	Tidak Najis"
            }
        }, {
            "id": 4,
            "soal": "4.	Cara Mengghilangkan Najis Sedang (mutawassithah) yaitu",
            "jawaban": {
                "a": "a.	Dengan mengucurinya dengan air atau mencucinya di dalam air yang banyak (direndam) hingga hilang rasa, bau dan warna dari bahan najisnya",
                "b": "b.	Di Diamkan Saja",
                "c": "c.	Di Bakar",
                "d": "d.	Dicuci tujuh kali dengan air dan salah satunya dengan tanah atau bahan lain yang mempunyai kemampuan menghilangkan rasa, bau dan warna"
            }
        }, {
            "id": 5,
            "soal": "5.	Cara Mengghilangkan Najis Berat (mughallazhah) yaitu",
            "jawaban": {
                "a": "a.	Dengan mengucurinya dengan air atau mencucinya di dalam air yang banyak (direndam) hingga hilang rasa, bau dan warna dari bahan najisnya.",
                "b": "b.	Di Diamkan Saja",
                "c": "c.	Di Bakar",
                "d": "d.	Dicuci tujuh kali dengan air dan salah satunya dengan tanah atau bahan lain yang mempunyai kemampuan menghilangkan rasa, bau dan warna."
            }
        }, {
            "id": 6,
            "soal": "6.	Cara menjaga Konsistensi dalam memproduksi Produk dan bahan yang halal perusahaan harus menerapkan",
            "jawaban": {
                "a": "a.	Sistem Keamanan Pangan",
                "b": "b.	Sistem Keselamatan Kerja",
                "c": "c.	Sistem Jaminan Halal",
                "d": "d.	Sistem Informasi"
            }
        }, {
            "id": 7,
            "soal": "7.	Kriteria Sistem Jaminan Halal Terdiri dari …. Kriteria",
            "jawaban": {
                "a": "a.	11",
                "b": "b.	20",
                "c": "c.	12",
                "d": "d.	14"
            }
        }, {
            "id": 8,
            "soal": "8.	Aktifitas manakah dibawah ini yang merupakan aktifitas kritis dalam Sistem Jaminan Halal",
            "jawaban": {
                "a": "a.	Seleksi Bahan Baru",
                "b": "b.	Pembelian",
                "c": "c.	Formulasi Produk baru",
                "d": "d.	Semua Benar"
            }
        }, {
            "id": 9,
            "soal": "9.	Setiap ada Bahan Baru perusahaan tidak wajib melaporkan kepada LPPOM MUI",
            "jawaban": {
                "a": "benar",
                "b": "salah"
            }
        }, {
            "id": 10,
            "soal": "10.	Audit Internal dilakukan 6 Bulan Sekali dan dilaporkan kepada LPPOM MUI",
            "jawaban": {
                "a": "benar",
                "b": "salah"
            }
        }
    ]
    try:
        
        header_style = doc.styles['Body Text']
        header_style.paragraph_format.space_after = 0
        header_style.font.bold = True
        header_style.font.size = Pt(11)
        
        task_style = doc.styles['List 2']
        task_style.paragraph_format.line_spacing = 1
        task_style.paragraph_format.space_before = Pt(2.8)
        task_style.paragraph_format.space_after = 0
        task_style.font.size = Pt(10)
        
        answer_style = doc.styles['List 3']
        answer_style.paragraph_format.line_spacing = 1
        answer_style.paragraph_format.space_after = 0
        answer_style.font.size = Pt(10)
        
        
        doc.add_paragraph(style=header_style).add_run('SOAL EVALUASI PELATIHAN INTERNAL SISTEM JAMINAN HALAL')
        sub_title = doc.add_paragraph(style=header_style).add_run('(soal ini diberikan kepada seluruh peserta pelatihan)')
        sub_title.font.italic = True
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-2].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(style=header_style).add_run(f'NAMA\t\t: {data_evaluasi["nama"]}')
        doc.add_paragraph(style=header_style).add_run(f'TANGGAL\t: {date.fromtimestamp(int(data_evaluasi["tanggal"]) / 1000).strftime("%d %B %Y")}')
        doc.add_paragraph(style=header_style).add_run(f'NILAI\t\t: ')
        doc.add_paragraph(style=header_style).add_run()
        
        for task in data:
            if "c" in task['jawaban']:
                doc.add_paragraph(style=task_style).add_run(task['soal'])
                for answer in task['jawaban']:
                    ans = doc.add_paragraph(style=answer_style).add_run(task['jawaban'][answer])
                    if data_evaluasi['data'][str(task['id'])] == answer:
                        ans.font.bold = True
            else:
                rad_task_style = doc.add_paragraph(style=task_style)
                rad_task_style.add_run(task['soal'])
                rad_task_ans_1 = rad_task_style.add_run("(Benar")
                rad_task_ans_1.font.bold = True
                rad_task_ans_2 = rad_task_style.add_run("/")
                rad_task_ans_2.font.bold = True
                rad_task_ans_3 = rad_task_style.add_run("Salah)")
                rad_task_ans_3.font.bold = True
                if data_evaluasi['data'][str(task['id'])] == "b":
                    rad_task_ans_1.font.strike = True
                else:
                    rad_task_ans_3.font.strike = True
        # doc.save('./app/assets/coba.docx')
        # return './app/assets/coba.docx'
    except Exception as error:
        traceback.print_exc()
        return error
    

def bukti_pelaksanan_docx(doc: Document, data_bukti_pelaksanan: dict, sign_data: dict = {}) -> str:
    """ function for generating Bukti pelaksanaan Pelatihan internal files """
    try:
        primary_style = doc.styles['Body Text']
        primary_style.font.size = Pt(11)
        primary_style.font.bold = True
        
        
        header = doc.add_paragraph(style=primary_style).add_run('Bukti Pelaksanaan Pelatihan Internal')
        header.font.size = Pt(14)
        header.font.bold = True
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        doc.add_paragraph(style=primary_style).add_run()
        doc.add_paragraph(style=primary_style).add_run(f'Tanggal Pelatihan\t: {date.fromtimestamp(int(data_bukti_pelaksanan["tanggal_pelaksanaan"]) / 1000).strftime("%d %B %Y")}')
        doc.add_paragraph(style=primary_style).add_run(f'Pemateri\t\t: {data_bukti_pelaksanan["pemateri"]}')
        doc.add_paragraph(style=primary_style).add_run('Materi\t\t\t: PENGENALAN HALAL HARAM DAN PENERAPAN SJH')
        
        
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[1].text = 'Nama Peserta'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[2].text = 'Posisi/Jabatan'
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[3].text = 'Tanda Tangan'
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[4].text = 'Nilai'
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        data = data_bukti_pelaksanan['data']
        for index in range(len(data)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[1].text = data[index]['nama']
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[2].text = data[index]['posisi']
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            if data[index]["ttd"]:
                row_cells[3].paragraphs[0].add_run().add_picture(f'./app/assets/{data[index]["ttd"]}', width=Pt(28))
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            else:
                row_cells[3].text = ""
            row_cells[4].text = str(data[index]['nilai'])
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        doc.add_paragraph(style=primary_style).add_run()
        
        notes = doc.add_paragraph(style=primary_style).add_run('Keterangan: \n1. Nilai Lulus  minimal 60, Jika tidak lulus harus melakukan remedial/pelatihan dan mengulang mengerjakan soal pelatihan \n2. Pelatihan internal harus menyertakan bukti foto (dokumentasi)')
        notes.font.size = Pt(10)
        
        doc.add_paragraph(style=primary_style).add_run()
        doc.add_paragraph(style=primary_style).add_run()
        
        if sign_data:
            detail_sign = doc.add_paragraph(style=primary_style).add_run(f'Jakarta, {date.fromtimestamp(int(data_bukti_pelaksanan["tanggal_pelaksanaan"]) / 1000).strftime("%d %B %Y")}\nMengetahui,')
            detail_sign.font.bold = False
            doc.add_paragraph(style=primary_style).add_run(f'{sign_data["title"]},')
            if sign_data['sign']:
                doc.add_picture(f'./app/assets/{sign_data["sign"]}', width=Inches(0.7))
            else:
                doc.add_paragraph(style=primary_style).add_run()
                doc.add_paragraph(style=primary_style).add_run()
                doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run(f'({sign_data["name"]})')
        else:
            detail_sign = doc.add_paragraph(style=primary_style).add_run('Jakarta, \nMengetahui,')
            detail_sign.font.bold = False
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run("Ttd + nama jelas")
        
        # doc.save('./app/assets/coba.docx')
        
        # return './app/assets/coba.docx'
    except Exception as error:
        traceback.print_exc()
        return error
    

def audit_internal_docx(doc: Document, jawaban_audit: dict, sign_data: dict = {}): 
    try:
        pertanyaan = [
            "Apakah kebijakan halal telah dijelaskan pada semua karyawan ?",
            "Apakah ada bukti sosialisasi kebijakan halal ? (daftar hadir sosialisasi)",
            "Apakah tersedia poster kebijakan halal dan edukasi halal di kantor, area produksi dan udang?",
            "Apakah ketua/anggota Tim Manajemen Halal telah mengikuti pelatihan eksternal setidaknya sekali dalam dua tahun?",
            "Apakah ada bukti pelatihan eksternal (sertifikat pelatihan) ?",
            "Apakah pelatihan internal kepada semua karyawan, termasuk karyawan baru, dengan materi seperti tercantum dalam Lampiran 3 telah dilaksanakan setidaknya setahun sekali ?",
            "Apakah ada bukti pelatihan internal (daftar hadir pelatihan) ?",
            "Apakah Daftar Bahan dengan format seperti pada Lampiran 4 telah dibuat ?",
            "Apakah nama/merk bahan dan nama produsen bahan yang dibeli sesuai dengan yang tercantum dalam Daftar Bahan Halal ?",
            "Apakah bukti pembelian (nota/kuitansi) dan contoh label kemasan (jika ada) selalu disimpan setidaknya selama 6 bulan ?",
            "Apakah setiap ada bahan baru selalu dimintakan persetujuan ke LPPOM MUI sebelum digunakan? (kecuali bahan tidak kritis dan bahan bersertifikat halal MUI yang ada di www.halalmui.org)",
            "Apakah bukti persetujuan penggunaan bahan baru dari LPPOM MUI selalu disimpan setidaknya selama dua tahun ?",
            "Apakah dilakukan pemeriksaan label bahan pada setiap pembelian atau penerimaan bahan ? (kecuali bahan tidak kritis)",
            "Apakah hasil pemeriksaan menunjukkan informasi nama bahan dan produsen yang tercantum di label sesuai dengan Daftar Bahan Halal ?",
            "Apakah ada formula/resep produk baku (untuk produk yang memiliki formula) ?",
            "Apakah bahan yang digunakan dalam produksi hanya bahan yang tercantum dalam Daftar Bahan ?",
            "Apakah formula produk yang digunakan pada proses produksi mengacu pada formula baku ?",
            "Jika terlanjur ada penggunaan bahan yang tidak tercantum dalam Daftar Bahan Halal, apakah produk yang dihasilkan tidak akan dijual ke konsumen dan dimusnahkan ?",
            "Apakah semua fasilitas dan peralatan produksi selalu dalam keadaan bersih (bebas dari najis) sebelum dan sesudah digunakan ?",
            "Apakah bahan dan produk selalu disimpan di tempat yang bersih dan terhindar dari najis?",
            "Apakah kendaraan yang digunakan untuk mengangkut produk halal dalam kondisi baik dan tidak digunakan untuk mengangkut produk lain yang diragukan kehalalannya ?",
            "Apakah setiap ada produk baru dengan merk yang sama selalu disertifikasi halal sebelum dipasarkan?",
            "Apakah setiap ada penambahan fasilitas produksi baru selalu didaftarkan untuk disertifikasi ?",
            "Apakah telah dilakukan audit internal setiap enam bulan sekali dengan cara memeriksa pelaksanaan seluruh prosedur operasional ? *",
            "Apakah audit internal dilakukan oleh ketua/ anggota Tim Manajemen Halal yang sudah mengikuti pelatihan ? *",
            "Apakah ada bukti pelaksanaan audit internal? *",
            "Apakah hasil audit internal telah dibahas dalam rapat kaji ulang manajemen yang dihadiri oleh ketua dan anggota Tim Manajemen Halal ? *",
            "Jika dalam audit internal ditemukan kelemahan, yaitu ada pertanyaan yang dijawab “tidak”, apakah segera dilakukan perbaikan agar kelemahan tersebut tidak terulang ? *",
            "Jika dalam audit internal ditemukan kelemahan, apakah ada bukti pelaksanaan perbaikan ? *",
            "Apakah ada bukti pelaksanaan rapat kaji ulang manajemen ? *",
            "Apakah form hasil audit internal yang telah terisi telah dikirimkan ke LPPOM MUI melalui Cerol? *"
        ]
        primary_style = doc.styles['Body Text']
        primary_style.font.size = Pt(11)
        
        header = doc.add_paragraph(style=primary_style).add_run('Lampiran 4.  Daftar Pertanyaan untuk Audit Internal')
        header.font.bold = True
        
        
        title = doc.add_paragraph(style=primary_style).add_run('AUDIT INTERNAL')
        title.font.size = Pt(14)
        title.font.bold = True
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        doc.add_paragraph(style=primary_style).add_run()
        doc.add_paragraph(style=primary_style).add_run(f'Tanggal Audit \t\t: {date.fromtimestamp(int(jawaban_audit["created_at"]) / 1000).strftime("%d %B %Y")}')
        doc.add_paragraph(style=primary_style).add_run(f'Auditor (yang mengaudit)\t: {jawaban_audit["auditee"]}')
        doc.add_paragraph(style=primary_style).add_run(f'Bagian yang di audit\t\t: {jawaban_audit["bagian_diaudit"]}')
        
        data = jawaban_audit['data']
        table = doc.add_table(rows=2, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'NO'
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(0,0).merge(table.cell(1,0))
        
        hdr_cells[1].text = 'PERTANYAAN'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(0,1).merge(table.cell(1,1))
        
        
        hdr_cells[2].text = 'HASIL AUDIT'
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(0,2).merge(table.cell(0,4))
        
        hdr_cells2 = table.rows[1].cells
        hdr_cells2[2].text = 'YA'
        hdr_cells2[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells2[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells2[3].text = 'TIDAK'
        hdr_cells2[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells2[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells2[4].text = 'KETERANGAN'
        hdr_cells2[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells2[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        for index in range(len(data)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[0].width = Inches(0.26)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[1].text = pertanyaan[index]
            row_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            if data[index]['jawaban']:
                row_cells[2].text = "v"
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
                row_cells[3].text = ""
            else:
                row_cells[2].text = ""
                row_cells[3].text = "x"
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            row_cells[4].text = str(data[index]['keterangan'])
            row_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    
        for cell in table.columns[1].cells:
            cell.width = Inches(3)
        for cell in table.columns[2].cells:
            cell.width = Inches(0.3)
        doc.add_paragraph(style=primary_style).add_run('*)Keterangan: Khusus pertanyaan mengenai audit internal dan kaji ulang manajemen, auditor dapat memeriksa pelaksanaan audit internal dan kaji ulang manajemen pada periode sebelumnya.')
        doc.add_paragraph(style=primary_style).add_run()
        
        if sign_data:
            doc.add_paragraph(style=primary_style).add_run(f'{sign_data["title"]},')
            if sign_data['sign']:
                doc.add_picture(f'./app/assets/{sign_data["sign"]}', width=Inches(0.7))
            else:
                doc.add_paragraph(style=primary_style).add_run()
                doc.add_paragraph(style=primary_style).add_run()
                doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run(f'({sign_data["name"]})')
        else:
            detail_sign = doc.add_paragraph(style=primary_style).add_run('Jakarta, \nMengetahui,')
            detail_sign.font.bold = False
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run("Ttd + nama jelas")
        
        # doc.save('./app/assets/coba.docx')
        
        # return './app/assets/coba.docx'
    except Exception as error:
        traceback.print_exc()
        return error
    
def data_hadir_kaji_ulang_docx(doc: Document(), daftar_hasil_kaji: dict, sign_data: dict):
    try:
        primary_style = doc.styles['Body Text']
        primary_style.font.size = Pt(11)
        
        
        header = doc.add_paragraph(style=primary_style).add_run('Lampiran 5.  Format Hasil Rapat Kaji Ulang Manajemen')
        header.font.bold = True
        
        title = doc.add_paragraph(style=primary_style).add_run('DAFTAR HADIR KAJI ULANG MANAJEMEN')
        title.font.size = Pt(14)
        title.font.bold = True
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        doc.add_paragraph(style=primary_style).add_run(f'Tanggal Kaji Ulang Manajemen\t :{date.fromtimestamp(int(daftar_hasil_kaji["tanggal"]) / 1000).strftime("%d %B %Y")}')
        # doc.add_paragraph(style=primary_style).add_run(f'Nama Pimpinan Perusahaan \t\t: {sign_data["name"]}')
        
        doc.add_paragraph(style=primary_style).add_run()
        
        
        list_orang = daftar_hasil_kaji['list_orang']
        pembahasan = daftar_hasil_kaji['pembahasan']
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[1].text = 'Nama'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[2].text = 'Jabatan/Bagian'
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[3].text = 'Paraf'
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        for index in range(len(list_orang)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[1].text = list_orang[index]['nama']
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[2].text = list_orang[index]['jabatan']
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[3].text = list_orang[index]['paraf']
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        doc.add_paragraph(style=primary_style).add_run()
        doc.add_paragraph(style=primary_style).add_run()
        table2 = doc.add_table(rows=1, cols=3)
        table2.style = 'Table Grid'
        hdr_cells = table2.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[1].text = 'Pembahasan Kaji Ulang'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[2].text = 'Perbaikan (jika ada)'
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for index in range(len(pembahasan)):
            row_cells = table2.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[1].text = pembahasan[index]['pembahasan']
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[2].text = pembahasan[index]['perbaikan']
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        doc.add_paragraph(style=primary_style).add_run()
        doc.add_paragraph(style=primary_style).add_run()
        
        if sign_data:
            doc.add_paragraph(style=primary_style).add_run(f'{sign_data["title"]},')
            if sign_data['sign']:
                doc.add_picture(f'./app/assets/{sign_data["sign"]}', width=Inches(0.7))
            else:
                doc.add_paragraph(style=primary_style).add_run()
                doc.add_paragraph(style=primary_style).add_run()
                doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run(f'({sign_data["name"]})')
        else:
            detail_sign = doc.add_paragraph(style=primary_style).add_run('Jakarta, \nMengetahui,')
            detail_sign.font.bold = False
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run("Ttd + nama jelas")

        # doc.save('./app/assets/coba.docx')
        
        # return './app/assets/coba.docx'
    except Exception as error:
        traceback.print_exc()
        return error


def surat_pernyataan_daftar_alamat_docx(doc: Document(), company, sign_data):
    try:
        normal_text = doc.styles['Body Text']
        normal_text.font.size = Pt(11)
        normal_text.font.bold = False


        header = doc.add_paragraph(style=normal_text).add_run("lampiran 6. Surat Pernyataan Daftar Alamat Fasilitas Produksi Dan Bebas Dari Babi dan Turunannya")
        header.font.bold = True

        doc.add_paragraph(style=normal_text).add_run()

        title = doc.add_paragraph(style=normal_text).add_run('SURAT PERNYATAAN DAFTAR ALAMAT FASILITAS PRODUKSI DAN BEBAS DARI BABI DAN TURUNANNYA')
        title.font.size = Pt(14)
        title.font.bold = True
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(style=normal_text).add_run()
        doc.add_paragraph(style=normal_text).add_run()

        doc.add_paragraph(style=normal_text).add_run("Saya yang bertanda tangan dibawah ini :")
        doc.add_paragraph(style=normal_text).add_run()
        doc.add_paragraph(style=normal_text).add_run()

        doc.add_paragraph(style=normal_text).add_run(f"Nama\t\t: {sign_data['name']}")
        # doc.add_paragraph(style=normal_text).add_run(f"No. Ktp\t\t: {company['ktp']}")
        doc.add_paragraph(style=normal_text).add_run(f"No. Ktp\t\t:")
        doc.add_paragraph(style=normal_text).add_run(f"No. Telepon\t: {company['company_number']}")
        doc.add_paragraph(style=normal_text).add_run(f"Jabatan\t\t: {sign_data['title']}")

        doc.add_paragraph(style=normal_text).add_run()
        doc.add_paragraph(style=normal_text).add_run()

        doc.add_paragraph(style=normal_text).add_run(f"Menyatakan bahwa, alamat produksi untuk perusahaan ({company['company_name']}) yaitu :")

        doc.add_paragraph(style=normal_text).add_run()

        doc.add_paragraph(style=normal_text).add_run(f"Seluruh fasilitas tersebut dan peralatan yang kami gunakan untuk produksi adalah bebas dari cemaran babi & turunannya.")
        doc.add_paragraph(style=normal_text).add_run(f"Demikian pernyataan ini saya buat dengan sebenar sebenarnya untuk dapat dipergunakan sebagaimana mestinya.")

        doc.add_paragraph(style=normal_text).add_run()

        detail_sign = doc.add_paragraph(style=normal_text).add_run(f'Jakarta, {date.today().strftime("%d %B %Y")}')
        detail_sign.font.bold = False
        if sign_data:
            doc.add_paragraph(style=normal_text).add_run(f'{sign_data["title"]},')
            if sign_data['sign']:
                doc.add_picture(f'./app/assets/{sign_data["sign"]}', width=Inches(0.7))
            else:
                doc.add_paragraph(style=normal_text).add_run()
                doc.add_paragraph(style=normal_text).add_run()
                doc.add_paragraph(style=normal_text).add_run()
            doc.add_paragraph(style=normal_text).add_run(f'({sign_data["name"]})')
        else:
            doc.add_paragraph(style=normal_text).add_run()
            doc.add_paragraph(style=normal_text).add_run()
            doc.add_paragraph(style=normal_text).add_run()
            doc.add_paragraph(style=normal_text).add_run("Ttd + nama jelas")

        # doc.save('./app/assets/coba.docx')
        
        # return './app/assets/coba.docx'
    except Exception as error:
        traceback.print_exc()
        return error

def form_pembelian_pemeriksaan_bahan_docx(doc: Document(),pembelian: list, pembelian_impor: list):
    try:
        
        primary_style = doc.styles['Body Text']
        primary_style.font.size = Pt(11)
        
        header = doc.add_paragraph(style=primary_style).add_run('Lampiran 7.  Format Form Aktivitas Kritis')
        header.font.bold = True
        
        title = doc.add_paragraph(style=primary_style).add_run('FORM PEMBELIAN DAN PEMERIKSAAN BAHAN')
        title.font.size = Pt(14)
        title.font.bold = True
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(style=primary_style).add_run()
        
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[1].text = 'Tanggal datang/beli'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[2].text = 'Nama/Merk Bahan'
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[3].text = 'Nama & Negara Produsen'
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[4].text = 'Ada di Daftar Bahan Halal? (Ya/Tidak)'
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[5].text = 'Exp. Date Bahan'
        hdr_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[5].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[6].text = 'Paraf'
        hdr_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for index in range(len(pembelian)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[1].text = date.fromtimestamp(int(pembelian[index]['Tanggal']) / 1000).strftime("%d %B %Y")
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[2].text = pembelian[index]['nama_dan_merk']
            row_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[3].text = pembelian[index]['nama_dan_negara']
            row_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[4].text = "Ya" if pembelian[index]['halal'] else "Tidak"
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[5].text = date.fromtimestamp(int(pembelian[index]['exp_bahan']) / 1000).strftime("%d %B %Y")
            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[5].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            if pembelian[index]['paraf']:
                row_cells[6].text = ""
                # row_cells[6].paragraphs[0].add_run().add_picture(f'./app/assets/{pembelian[index]["paraf"]}', width=Pt(28))
                # row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # row_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            else:
                row_cells[6].text = ""
        
        doc.add_paragraph(style=primary_style).add_run()
        sub_head = doc.add_paragraph(style=primary_style).add_run('Khusus Daging Impor')
        sub_head.font.bold = True
        
        table2 = doc.add_table(rows=1, cols=7)
        table2.style = 'Table Grid'
        hdr_cells = table2.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[1].text = 'Tanggal datang/beli'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[2].text = 'Nama/Merk/Kode Bahan'
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[3].text = 'Nama & Lokasi Produsen'
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[4].text = 'Ada di Daftar Bahan Halal? (Ya/Tidak)'
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[5].text = 'Kesesuaian nomor lot dengan SH per pengapalan (Ya/Tdk)'
        hdr_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[5].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[6].text = 'Paraf'
        hdr_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for index in range(len(pembelian_impor)):
            row_cells = table2.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[1].text = date.fromtimestamp(int(pembelian_impor[index]['Tanggal']) / 1000).strftime("%d %B %Y")
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[2].text = pembelian_impor[index]['nama_dan_merk']
            row_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[3].text = pembelian_impor[index]['nama_dan_negara']
            row_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[4].text = "Ya" if pembelian_impor[index]['halal'] else "Tidak"
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[5].text = date.fromtimestamp(int(pembelian_impor[index]['exp_bahan']) / 1000).strftime("%d %B %Y")
            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[5].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            if pembelian_impor[index]['paraf']:
                row_cells[6].text = ""
                # row_cells[6].paragraphs[0].add_run().add_picture(f'./app/assets/{pembelian_impor[index]["paraf"]}', width=Pt(28))
                # row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # row_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            else:
                row_cells[6].text = ""
        
        doc.add_paragraph(style=primary_style).add_run()
        keterangan = doc.add_paragraph(style=primary_style).add_run('Sesuaikan dengan yang digunakan perusahaan saat ini untuk masing-masing aktivitas kritis, diatas hanya contoh saja')
        keterangan.font.italic = True
        
        # doc.save('./app/assets/coba.docx')
        
        # return './app/assets/coba.docx'
    except Exception as error:
        traceback.print_exc()
        return error
    
def form_stok_bahan_docx(doc: Document, stok_barang: list):
    try:
        
        primary_style = doc.styles['Body Text']
        primary_style.font.size = Pt(11)
        
        
        header = doc.add_paragraph(style=primary_style).add_run('FORM STOK BAHAN')
        header.font.size = Pt(14)
        header.font.bold = True
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(style=primary_style).add_run()
        
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[1].text = 'Tanggal Pembelian'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[2].text = 'Nama Bahan'
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[3].text = 'Jumlah Bahan Masuk'
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[4].text = 'Jumlah Bahan Keluar'
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[5].text = 'Sisa Stok'
        hdr_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[5].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[6].text = 'Paraf'
        hdr_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for index in range(len(stok_barang)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[1].text = date.fromtimestamp(int(stok_barang[index]['tanggal_beli']) / 1000).strftime("%d %B %Y")
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[2].text = stok_barang[index]['nama_bahan']
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[3].text = stok_barang[index]['jumlah_bahan']
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[4].text = stok_barang[index]['jumlah_keluar']
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[5].text = stok_barang[index]['stok_sisa']
            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[5].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
            if stok_barang[index]['paraf']:
                row_cells[6].text = ""
                # row_cells[6].paragraphs[0].add_run().add_picture(f'./app/assets/{pembelian_impor[index]["paraf"]}', width=Pt(28))
                # row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # row_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            else:
                row_cells[6].text = ""
        
        # doc.save('./app/assets/coba.docx')
        
        # return './app/assets/coba.docx'
    except Exception as error:
        traceback.print_exc()
        return error
    
    
def form_produksi_docx(doc: Document ,produksi: list):
    try:
        
        primary_style = doc.styles['Body Text']
        primary_style.font.size = Pt(11)
        
        
        header = doc.add_paragraph(style=primary_style).add_run('FORM PRODUKSI')
        header.font.size = Pt(14)
        header.font.bold = True
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(style=primary_style).add_run()
        
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[1].text = 'Tanggal produksi'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[2].text = 'Nama Produk'
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[3].text = 'Jumlah Awal'
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[4].text = 'Jumlah Produk Keluar'
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[5].text = 'Sisa Stok'
        hdr_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[5].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[6].text = 'Paraf'
        hdr_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for index in range(len(produksi)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[1].text = date.fromtimestamp(int(produksi[index]['tanggal_produksi']) / 1000).strftime("%d %B %Y")
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[2].text = produksi[index]['nama_produk']
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[3].text = produksi[index]['jumlah_awal']
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[4].text = produksi[index]['jumlah_produk_keluar']
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[5].text = produksi[index]['sisa_stok']
            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[5].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
            if produksi[index]['paraf']:
                row_cells[6].text = ""
                # row_cells[6].paragraphs[0].add_run().add_picture(f'./app/assets/{pembelian_impor[index]["paraf"]}', width=Pt(28))
                # row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # row_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            else:
                row_cells[6].text = ""
        
        # doc.save('./app/assets/coba.docx')
        
        # return './app/assets/coba.docx'
    except Exception as error:
        traceback.print_exc()
        return error
    
    
def form_pemusnahan_barang_docx(doc: Document, kebersihan: list):
    try:
        
        primary_style = doc.styles['Body Text']
        primary_style.font.size = Pt(11)
        
        
        header = doc.add_paragraph(style=primary_style).add_run('LAPORAN PEMUSNAHAN BARANG/PRODUK')
        header.font.size = Pt(14)
        header.font.bold = True
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(style=primary_style).add_run()
        
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[1].text = 'Nama Produk'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[2].text = 'Kode/Tgl Produksi'
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[3].text = 'Jumlah'
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[4].text = 'Penyebab'
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[5].text = 'Tgl Pemusnahan'
        hdr_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[5].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[6].text = 'Penanggungjawab'
        hdr_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for index in range(len(kebersihan)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[1].text = date.fromtimestamp(int(kebersihan[index]['tanggal_produksi']) / 1000).strftime("%d %B %Y")
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[2].text = kebersihan[index]['nama_produk']
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[3].text = kebersihan[index]['jumlah']
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[4].text = kebersihan[index]['penyebab']
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[5].text = date.fromtimestamp(int(kebersihan[index]['tanggal_pemusnahan']) / 1000).strftime("%d %B %Y")
            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[5].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[6].text = kebersihan[index]['penanggungjawab']
            row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        # doc.save('./app/assets/coba.docx')
        
        # return './app/assets/coba.docx'
    except Exception as error:
        traceback.print_exc()
        return error
    
    
def form_pengecheckan_kebersihan_docx(doc: Document, kebersihan: list):
    try:
        
        primary_style = doc.styles['Body Text']
        primary_style.font.size = Pt(11)
        
        
        header = doc.add_paragraph(style=primary_style).add_run('FORM PENGECEKAN KEBERSIHAN FASILITAS PRODUKSI DAN KENDARAAN')
        header.font.size = Pt(14)
        header.font.bold = True
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(style=primary_style).add_run()
        
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[1].text = 'Hari/Tgl'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[2].text = 'R. Produksi'
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[3].text = 'R. Gudang'
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[4].text = 'Peralatan dan Mesin Produksi'
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[5].text = 'Kebersihan Kendaraan'
        hdr_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[5].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[6].text = 'Penanggungjawab'
        hdr_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for index in range(len(kebersihan)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[1].text = date.fromtimestamp(int(kebersihan[index]['tanggal']) / 1000).strftime("%d %B %Y")
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[2].text = "v" if kebersihan[index]['produksi'] else "x"
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[3].text = "v" if kebersihan[index]['gudang'] else "x"
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[4].text = "v" if kebersihan[index]['mesin'] else "x"
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[5].text = "v" if kebersihan[index]['kendaraan'] else "x"
            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[5].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[6].text = kebersihan[index]['penanggungjawab']
            row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # doc.save('./app/assets/coba.docx')
        
        # return './app/assets/coba.docx'
    except Exception as error:
        traceback.print_exc()
        return error
    
    
def daftar_bahan_halal_docx(doc: Document, bahan_halal: list, company: dict, sign_data: dict):
    try:
        primary_style = doc.styles['Body Text']
        primary_style.font.size = Pt(11)
        
        header = doc.add_paragraph(style=primary_style).add_run('Lampiran 8.  Daftar Bahan Halal (TIDAK WAJIB DIISI)')
        header.font.bold = True
        
        title = doc.add_paragraph(style=primary_style).add_run('DAFTAR BAHAN HALAL')
        title.font.size = Pt(14)
        title.font.bold = True
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(style=primary_style).add_run(f'Nama Perusahaan\t :{company["company_name"]}')
        doc.add_paragraph(style=primary_style).add_run(f'Kelompok produk\t :{company["product_type"]}')
        
        doc.add_paragraph(style=primary_style).add_run()
        
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[1].text = 'Nama/Merk/ Kode Bahan'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[2].text = 'Nama dan Negara Produsen'
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[3].text = 'Pemasok'
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[4].text = 'Lembaga Penerbit'
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[5].text = 'Nomor'
        hdr_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[5].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[6].text = 'Masa Berlaku'
        hdr_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[7].text = 'Dokumen Lain'
        hdr_cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[7].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[8].text = 'Keterangan'
        hdr_cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[8].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for index in range(len(bahan_halal)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[1].text = bahan_halal[index]['nama_merk']
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[2].text = bahan_halal[index]['nama_negara']
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[3].text = bahan_halal[index]['pemasok']
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[4].text = bahan_halal[index]['penerbit']
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[4].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[5].text = bahan_halal[index]['nomor']
            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[5].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[6].text = bahan_halal[index]['masa_berlaku']
            row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[6].text = bahan_halal[index]['keterangan']
            row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[6].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        
        doc.add_paragraph(style=primary_style).add_run()
        
        if sign_data:
            doc.add_paragraph(style=primary_style).add_run(f'{sign_data["title"]},')
            if sign_data['sign']:
                doc.add_picture(f'./app/assets/{sign_data["sign"]}', width=Inches(0.7))
            else:
                doc.add_paragraph(style=primary_style).add_run()
                doc.add_paragraph(style=primary_style).add_run()
                doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run(f'({sign_data["name"]})')
        else:
            detail_sign = doc.add_paragraph(style=primary_style).add_run('Jakarta, \nMengetahui,')
            detail_sign.font.bold = False
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run("Ttd + nama jelas")

        
        # doc.save('./app/assets/coba.docx')
        
        # return './app/assets/coba.docx'
    except Exception as error:
        traceback.print_exc()
        return error
    
def matrik_produk_docx(doc: Document, matrik: list, company: dict, sign_data: dict):
    try:
        # doc = Document()
        primary_style = doc.styles['Body Text']
        primary_style.font.size = Pt(11)
        
        header = doc.add_paragraph(style=primary_style).add_run('Lampiran 9.  Matriks Produk  (TIDAK WAJIB DIISI)')
        header.font.bold = True
        
        title = doc.add_paragraph(style=primary_style).add_run('MATRIKS PRODUK')
        title.font.size = Pt(14)
        title.font.bold = True
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(style=primary_style).add_run(f'Nama Perusahaan\t :{company["company_name"]}')
        doc.add_paragraph(style=primary_style).add_run(f'Kelompok produk\t :{company["product_type"]}')
        
        doc.add_paragraph(style=primary_style).add_run()
        
        table = doc.add_table(rows=1, cols=(2+len(matrik)))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells[1].text = 'Nama Bahan'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        
        list_barang = []
        for bahan in matrik:
            for barang in bahan['list_barang']:
                list_barang.append(barang['barang'])
        if len(list_barang) > 1:
            list_barang = list(dict.fromkeys(list_barang))
        
        for index_barang in range(len(list_barang)):
            hdr_cells[2+index_barang].text = list_barang[index_barang]
            hdr_cells[2+index_barang].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[2+index_barang].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
        for index in range(len(matrik)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            row_cells[1].text = matrik[index]['nama_bahan']
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for index_barang in range(len(list_barang)):
                row_cells[2+index_barang].text = "v" if matrik[index]['list_barang'][index_barang]['status'] else "x"
                row_cells[2+index_barang].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[2+index_barang].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            
        
        if sign_data:
            doc.add_paragraph(style=primary_style).add_run(f'{sign_data["title"]},')
            if sign_data['sign']:
                doc.add_picture(f'./app/assets/{sign_data["sign"]}', width=Inches(0.7))
            else:
                doc.add_paragraph(style=primary_style).add_run()
                doc.add_paragraph(style=primary_style).add_run()
                doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run(f'({sign_data["name"]})')
        else:
            detail_sign = doc.add_paragraph(style=primary_style).add_run('Jakarta, \nMengetahui,')
            detail_sign.font.bold = False
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run()
            doc.add_paragraph(style=primary_style).add_run("Ttd + nama jelas")
        # doc.save('./app/assets/coba.docx')
        
        # return './app/assets/coba.docx'
    except Exception as error:
        traceback.print_exc()
        return error
    
def add_linespace(doc: Document):
    linespace_style = doc.styles['Body Text']
    linespace = doc.add_paragraph(style=linespace_style).add_run()
    linespace.font.size = 10

def add_head(doc: Document, text: str, size: int = 14, bold: bool = True):
    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{text}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.font.size = Pt(size)
    head.font.bold = bold 