# modules
from docx import Document
from docx.opc.coreprops import CoreProperties
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt

def add_linespace():
    linespace_style = doc.styles['Body Text']
    linespace = doc.add_paragraph(style=linespace_style).add_run()
    linespace.font.size = 10

def add_head(text: str, size: int = 14):
    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{text}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.font.size = Pt(size)
    head.font.bold = True 


# ? Page 1
pedoman_head = "PEDOMAN PENGISIAN TEMPLATE MANUAL SJH"
pedoman_body = """
Template Manual SJH merupakan format Manual SJH yang dibuat LPPOM MUI khusus untuk perusahaan kecil. Perusahaan diharapkan dapat memahami isi template Manual SJH agar dapat diterapkan di perusahaan. Perhatikan huruf yang berwarna merah agar disesuaikan dengan kondisi di perusahaan.
Berikut adalah hal yang perlu diperhatikan dalam penyusunan Manual SJH:"""
pedoman_list = [{
    "parent": "Informasi Umum Perusahaan pada Pendahuluan",
    "child": [
        "a.	Isi data sesuai dengan perusahaan yang mengajukan sertifikasi halal.", 
        "b.	Jika lokasi pabrik atau fasilitas (dapur/outlet restoran) berbeda dengan lokasi perusahaan, dapat ditambahkan sesuai dengan alamat pabrik dan dapur/outlet."
        ]
},
{
    "parent": "Kriteria Sistem Jaminan Halal:",
    "child": [
        "a.	Kebijakan Halal: Isi “….........., ..........................” dengan tempat dan tanggal/bulan/tahun sesuai dengan waktu penetapan Kebijakan Halal. Contoh : Bogor, 5 Januari 2014",
        "b.	Tim Manajemen Halal: disesuaikan dengan kondisi perusahaan. 1 orang tim halal bisa merangkap beberapa posisi.",
        "c.	Fasilitas produksi: disesuaikan dengan apa saja fasilitas yang digunakan pada produksi halal.",
        "d.	Prosedur aktivitas kritis: disesuaikan dengan kondisi perusahaan. Persetujuan bahan baru di LPPOM MUI Provinsi disesuaikan dengan provinsi dimana dilakukan pendaftaran sertifikasi halal.",
        "e.	Laporan berkala audit internal di LPPOM MUI Provinsi disesuaikan dengan provinsi dimana dilakukan pendaftaran sertifikasi halal."
    ]
},
{
    "parent": "Lampiran Manual SJH:",
    "child": [
        "a.	Ada 2 jenis lampiran: (i) Lampiran default dari template Manual SJH sebagai bahan rujukan tetapi dapat dimodifikasi oleh perusahaan; (ii) Lampiran tambahan jika ada dokumen yang ingin dilampirkan oleh perusahaan, misalnya Daftar Bahan, jadwal pelatihan, jadwal audit internal, jadwal kaji ulang manajemen, daftar bahan baru.",
        "b.	Lampiran 1: perusahaan dapat memodifikasi isi poster sesuai dengan kebutuhan perusahaan.",
        "c.	Lampiran 2: isi nama lengkap, jabatan di perusahaan serta ditandatangani. Jabatan dapat dirangkap, misalnya 1 orang merangkap sebagai Ketua Tim, R&D dan Pembelian.",
        "d.	Lampiran 3: tidak ada yang perlu diisi. Perusahaan dapat melengkapi materi pelatihan dengan pengetahuan SJH lain yang dianggap perlu untuk diketahui oleh karyawan.",
        "e.	Lampiran 4: isi data bahan yang digunakan untuk seluruh produk yang disertifikasi.",
        "f.	Lampiran 5: isi sesuai formula/resep produk.",
        "g.	Lampiran 6: isi form pemeriksaan bahan setelah bahan dibeli atau bahan diterima dari supplier. Pastikan bahan sesuai dengan yang terdapat di Daftar Bahan Halal. Jika tidak sesuai, maka lakukan evaluasi bahan baru mengikuti prosedur Pemilihan Bahan Baru.",
        "h.	Lampiran 7: daftar pertanyaan pada saat audit internal. Perusahaan dapat memodifikasi daftar pertanyaan sesuai dengan kebutuhan perusahaan.",
        "i.	Lampiran 8: isi form sesuai dengan pembahasan pada rapat kaji ulang manajemen."
    ]
}
]
doc =  Document()
sections = doc.sections
for section in sections:
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.26)
    section.top_margin = Inches(1.06)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.8)

section = doc.sections[0]

# ? ============================================================
# ! Header 
add_head(pedoman_head, 11)


body_style = doc.styles['Body Text']
body = doc.add_paragraph(style=body_style).add_run(f'{pedoman_body}')
body.font.size = Pt(10)


for _list in pedoman_list:
    body = doc.add_paragraph(style='List Number').add_run(f'{_list["parent"]}')
    body.font.size = Pt(10)
    for grand_child in _list["child"]:
        body = doc.add_paragraph(style='List 2').add_run(f'{grand_child}')
        body.font.size = Pt(10)


# * ============================================================
doc.add_page_break()
# * ============================================================

pendahuluan_head = "PENDAHULUAN"
pendahuluan_body = [{
    "parent": "I. Informasi Umum Perusahaan",
    "child": [
        "Nama Perusahaan \t: Faisal Outlet",
        "Alamat Perusahaan \t: jl spg 7 gang kober rt04 09 no 20 lubang buaya jakarta timur",
        "Nama Pabrik \t\t: Faisal Outlet",
        "Alamat Pabrik  \t\t: jl spg 7 gang kober rt04 09 no 20 lubang buaya jakarta timur",
        "Telp/Fax Perusahaan \t: 0878-3687-1856 ",
        "Contact Person/Email \t: faisalismail277@gmail.com",
        "Nama/Merk Produk \t: faisal outlet",
        "Jenis Produk \t\t: tahu baso faisal outlet / sosis solo frozen",
        "Daerah Pemasaran \t: provinsi / nasional / internasional",
        "Sistem Pemasaran \t: retail / non retail "
    ]
},{
    "parent": "II. Tujuan",
    "child": ["Manual SJH disusun untuk menjadi pedoman dalam penerapan SJH di perusahaan, dalam rangka menjaga kesinambungan produksi halal sesuai dengan persyaratan sertifikasi halal MUI."]
}]
# ? ============================================================
# ! Header 
add_head(pendahuluan_head)

for data in pendahuluan_body:
    body_style = doc.styles['Body Text']
    body = doc.add_paragraph(style='List').add_run(f'{data["parent"]}')
    body.font.size = Pt(11)
    body.font.bold = True
    for grand_child in data["child"]:
        body = doc.add_paragraph(style='List 2').add_run(f'{grand_child}')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body.font.size = Pt(10)
    doc.add_paragraph(style='Body Text').add_run()

kriteria_head = "KRITERIA SISTEM JAMINAN HALAL"
kebijakan_halal_head = "I.\tKebijakan Halal"
kebijakan_halal_card = {
    "header": "KEBIJAKAN HALAL",
    "company_name": "FAISAL OUTLET",
    "body": '“Kami berkomitmen tinggi untuk menghasilkan produk halal, dengan hanya menggunakan bahan yang telah disetujui oleh LPPOM MUI dan diproduksi dengan menggunakan peralatan yang bebas dari najis. Kami akan mencapainya dengan membentuk tim manajemen halal dan melaksanakan dengan sungguh-sungguh Sistem Jaminan Halal”',
    "date": "Jakarta, 12 Agustus 2021",
    "signature_title": "Pimpinan Perusahaan,",
    "signature_image": "mantap.png",
    "signature_name": "( Faisal Ismail )"
}
kebijakan_halal_body = "Kebijakan halal disosialisasikan ke seluruh karyawan dengan menempel Kebijakan halal dan poster halal (Lampiran 1) di lokasi yang mudah dilihat seperti di area kantor dan produksi, serta melalui email ke supplier bahan baku (jika ada supplier)."

# ? ============================================================
# ! Header 
add_head(kriteria_head)

add_linespace()

body = doc.add_paragraph(style='List').add_run(f'{kebijakan_halal_head}')
body.font.size = Pt(11)
body.font.bold = True

heading_style = doc.styles['Body Text']
head=doc.add_paragraph(style=heading_style).add_run(f'{kebijakan_halal_card["header"]}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
head.font.size = Pt(11)
head.font.bold = True 

heading_style = doc.styles['Body Text']
head=doc.add_paragraph(style=heading_style).add_run(f'{kebijakan_halal_card["company_name"]}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
head.font.size = Pt(11)
head.font.bold = True 


heading_style = doc.styles['Body Text']
head=doc.add_paragraph(style=heading_style).add_run(f'{kebijakan_halal_card["body"]}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
head.font.size = Pt(11)

add_linespace()

heading_style = doc.styles['Body Text']
head=doc.add_paragraph(style=heading_style).add_run(f'{kebijakan_halal_card["date"]}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
head.font.size = Pt(11)

heading_style = doc.styles['Body Text']
head=doc.add_paragraph(style=heading_style).add_run(f'{kebijakan_halal_card["signature_title"]}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
head.font.size = Pt(11)

# signature image & text
ceo_sign = doc.styles['Body Text']
doc.add_picture(kebijakan_halal_card["signature_image"], width=Inches(1))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(style=ceo_sign).add_run(f'{kebijakan_halal_card["signature_name"]}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
ceo_sign.font.size = Pt(11)

add_linespace()

heading_style = doc.styles['Body Text']
head=doc.add_paragraph(style=heading_style).add_run(f'{kebijakan_halal_body}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
head.font.size = Pt(11)


# * ============================================================
doc.add_page_break()
# * ============================================================
tim_manajemen_head = "II.\tTim Manajemen Halal "
tim_manajemen_body = "Untuk menerapkan SJH dan dalam rangka menjaga konsistensi kehalalan produk, dengan ini ditunjuk Tim Manajemen Halal yang terdiri dari pimpinan perusahaan dan karyawan seperti tercantum dalam Lampiran 2.  Tugas Tim Manajemen Halal sebagai berikut :"

tim_manajemen_table = (
    ("Ketua Tim/Pimpinan", (
        "1. Bertanggungjawab dalam sosialisasi kebijakan halal",
        "2. Bertanggungjawab dalam penunjukkan tim manajemen halal",
        "3. Bertanggungjawab dalam pelatihan eksternal/internal",
        "4. Bertanggungjawab dalam audit internal dan pengiriman laporan berkala ke LPPOM MUI",
        "5. Mengkoordinasikan kaji ulang manajemen"
        )
    ),
    ("Pembelian", (
        "1. Bertanggungjawab dalam proses seleksi bahan baru",
        "2. Bertanggungjawab dalam proses pembelian bahan baku",
        "3. Membuat dan memperbaharui Daftar Bahan Halal serta memonitoring masa berlaku dokumen bahan",
        "4. Bertanggungjawab dalam proses pemeriksaan kedatangan bahan",
        ),
    ),
    ("Produksi", (
        "1. Bertanggungjawab dalam proses produksi halal",
        "2. Bertanggungjawab dalam proses pengembangan produk",
        "3. Bertanggungjawab dalam proses pencucian fasilitas produksi",
        "4. Bertanggungjawab dalam pengananan produk yang tidak memenuhi kriteria (jika terjadi)",
        ),
    ),
    ("Penyimpanan dan Pengiriman", (
        "1. Bertanggungjawab dalam proses penyimpanan bahan dan produk jadi",
        "2. Bertanggungjawab dalam proses transportasi bahan dan produk jadi"
    ))
)
tim_manajemen_note = "*) Jabatan dapat dirangkap, misalnya 1 orang merangkap sebagai Ketua Tim, R&D dan Pembelian"

body = doc.add_paragraph(style='List').add_run(f'{tim_manajemen_head}')
body.font.size = Pt(11)
body.font.bold = True

body = doc.add_paragraph(style='List').add_run(f'\t{tim_manajemen_body}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
body.font.size = Pt(11)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Bagian/Jabatan'
hdr_cells[1].text = 'Tugas Tim Manajemen Halal'
for bagian, tugas_tugas in tim_manajemen_table:
    row_cells = table.add_row().cells
    row_cells[0].text = bagian
    for tugas in tugas_tugas:
        row_cells[1].text += tugas + "\n"
table.allow_autofit = True
for cell in table.columns[0].cells:
    cell.width = Inches(1)
for cell in table.columns[1].cells:
    cell.width = Inches(5.56)

body = doc.add_paragraph(style='List 2').add_run(f'{tim_manajemen_note}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
body.font.size = Pt(11)

add_linespace()

pelatihan_head = "III.	Pelatihan"
pelatihan_list = [
    "a.	Tujuan pelatihan yaitu meningkatkan pemahaman Tim Manajemen Halal tentang persyaratan sertifikasi halal.",
    "b.	Pelatihan dilakukan secara internal dan eksternal.",
    "c.	Pelatihan internal dilakukan 1 tahun sekali oleh personel yang sudah mengikuti pelatihan halal. Pelatihan diikuti oleh seluruh karyawan. Setiap karyawan baru harus mendapatkan pelatihan internal sebelum mulai bekerja. Materi dan soal tes pelatihan internal sesuai dengan Lampiran 3 atau dapat dimodifikasi sesuai kebutuhan. Daftar hadir pelatihan dan hasil tes disimpan selama dua tahun sebagai bukti pelatihan internal.",
    "d.	Salah satu tim manajemen halal akan mengikuti pelatihan eksternal ke lembaga pelatihan halal dengan program pelatihan berbasis HAS 23000. Sertifikat pelatihan eksternal disimpan selama dua tahun sebagai bukti pelatihan eksternal. Pelatihan eksternal dapat diikuti setiap 2 tahun sekali atau sesuai kebutuhan perusahaan."
]

body = doc.add_paragraph(style='List').add_run(f'{pelatihan_head}')
body.font.size = Pt(11)
body.font.bold = True

for pelatihan in pelatihan_list:
    body = doc.add_paragraph(style='List 2').add_run(f'{pelatihan}')
    body.font.size = Pt(11)

add_linespace()

bahan_head = "IV.	Bahan"
bahan_list = [
    "a.	Perusahaan hanya menggunakan bahan yang sesuai dengan kriteria SJH dan disetujui oleh LPPOM MUI untuk menghasilkan produk yang disertifikasi.",
    "b.	Perusahaan membuat Daftar Bahan Halal yaitu daftar bahan yang disetujui LPPOM MUI, yang dapat dilihat pada Lampiran 4.",
    "c.	Daftar Bahan Halal digunakan sebagai acuan dalam proses pembelian, pemeriksaan bahan datang dan proses produksi.",
    "d.	Perbaikan Daftar Bahan Halal dilakukan bila terjadi perubahan bahan.",
    "e.	Semua bahan dilengkapi dengan dokumen pendukung, kecuali bahan tidak kritis (positive list) yang tercantum dalam SK LPPOM MUI. Dokumen pendukung dapat berupa Sertifikat halal atau dokumen lain dari produsen bahan yang menjelaskan sumber bahan.",
    "f.	Ketua tim manajemen halal atau wakilnya akan memeriksa masa berlaku Sertifikat halal yang terdapat di Daftar Bahan Halal secara berkala. Jika masa berlaku Sertifikat halal akan berakhir maka dilakukan pemeriksaan masa berlaku Sertifikat halal di website LPPOM MUI (www.halalmui.org) atau meminta Sertifikat halal terbaru ke supplier. Data Sertifikat halal terbaru akan dimasukkan ke dalam Daftar Bahan Halal."
]

body = doc.add_paragraph(style='List').add_run(f'{bahan_head}')
body.font.size = Pt(11)
body.font.bold = True

for bahan in bahan_list:
    body = doc.add_paragraph(style='List 2').add_run(f'{bahan}')
    body.font.size = Pt(11)

add_linespace()

produk_head = "V.	Produk"
produk_list = [
    "a.	Perusahaan hanya akan memproduksi produk halal dengan nama produk yang tidak menggunakan nama yang mengarah pada sesuatu yang diharamkan atau ibadah yang tidak sesuai dengan syariah Islam, bau/rasa produk tidak mengarah kepada produk haram, bentuk produk tidak berupa anjing atau babi, serta kemasan produk tidak vulgar.",
    "b.	Perusahaan akan membuat Matriks Produk untuk semua produk yang disertifikasi halal sesuai dengan formula/resep pada proses produksi, yang dapat dilihat pada Lampiran 5.",
    "c.	Perusahaan akan mendaftarkan seluruh Produk pangan eceran (retail) dengan merk sama yang beredar di Indonesia. Untuk produk pangan bukan eceran (non retail) yang mempunyai merk/brand dan hanya didaftarkan sebagian, maka kami akan mencantumkan logo halal MUI untuk produk yang disertifikasi."
]

body = doc.add_paragraph(style='List').add_run(f'{produk_head}')
body.font.size = Pt(11)
body.font.bold = True

for produk in produk_list:
    body = doc.add_paragraph(style='List 2').add_run(f'{produk}')
    body.font.size = Pt(11)

add_linespace()

fasilitas_produksi_head = "VI.	Fasilitas Produksi"
fasilitas_produksi_list = [
    "a.	Fasilitas produksi harus bersih dan terhindar dari kontaminasi bahan najis.",
    "b.	Fasilitas produksi mencakup seluruh fasilitas yang digunakan untuk proses produksi seperti gudang penyimpanan, ruang pencucian, ruang produksi, dapur/outlet.",
    "c.	Fasilitas produksi yang kami gunakan untuk menghasilkan produk halal tidak digunakan untuk menghasilkan produk lain yang mengandung bahan babi.",
    "d.	Setiap ada tambahan fasilitas produksi yang digunakan untuk menghasilkan produk, baik milik perusahaan sendiri maupun milik pihak lain, maka akan didaftarkan untuk disertifikasi halal dan menjadi ruang lingkup implementasi Sistem Jaminan Halal."
]

body = doc.add_paragraph(style='List').add_run(f'{fasilitas_produksi_head}')
body.font.size = Pt(11)
body.font.bold = True

for fasilitas_produksi in fasilitas_produksi_list:
    body = doc.add_paragraph(style='List 2').add_run(f'{fasilitas_produksi}')
    body.font.size = Pt(11)

add_linespace()


prosedur_aktifitas_head = "VII.	Prosedur Aktifitas Kritis"
prosedur_aktifitas_body = "Berikut adalah prosedur sederhana yang terkait dengan halal :"
prosedur_aktifitas_table = (
    ("Pemilihan Bahan Baru", (
        "1. Jika produksi produk halal akan menggunakan bahan baru di luar Daftar Bahan Halal (termasuk bahan lama dengan produsen baru), maka kami akan mencari bahan positif list atau bahan bersertifikat halal MUI yang ada di www.halalmui.org.",
        "2. Jika bahan bukan termasuk kategori diatas, maka harus meminta dokumen bahan ke produsen (dapat berupa pernyataan, spesifikasi, diagram alir proses) atau Sertifikat Halal dari lembaga yang diakui MUI, kemudian meminta persetujuan bahan tersebut ke LPPOM MUI.",
        "3. Pengajuan persetujuan ke LPPOM MUI dilakukan melalui Cerol (menu Inquiry Material Approval) atau melalui email (sesuai dengan email LPPOM MUI Provinsi). Bukti persetujuan bahan baru disimpan setidaknya selama dua tahun.",
        "4. Bahan baru dimasukkan ke dalam Daftar Bahan Halal dan disebarkan kepada seluruh tim halal. Bahan baru dapat digunakan setelah mendapatkan persetujuan tertulis dari LPPOM MUI."
    )),
    ("Pembelian", (
        "1.	Melakukan pembelian bahan dengan nama/merk dan produsen sesuai dengan yang tercantum dalam Daftar Bahan Halal.",
        "2.	Bukti pembelian (nota/kuitansi) disimpan setidaknya selama 6 bulan, kecuali untuk bahan yang jarang dibeli maka disimpan bukti pembelian terakhir."
    )),
    ("Pemeriksaan Bahan Datang", (
        "1.	Memeriksa label bahan pada setiap pembelian atau penerimaan bahan untuk memastikan kesesuaian nama bahan, nama produsen dan negara produsen dengan yang tercantum dalam Daftar Bahan Halal. Bahan yang boleh digunakan hanya bahan yang namanya, nama produsen dan negara produsennya sesuai dengan Daftar Bahan Halal.",
        "2.	Pemeriksaan label bahan dituliskan dalam form pemeriksaan bahan seperti pada Lampiran 6.",
        "3.	Pengecualian untuk bahan positif list tidak dilakukan pemeriksaan label bahan."
    )), 
    ("Produksi", (
        "1.	Membuat formula/resep produk yang akan menjadi acuan/rujukan untuk bagian produksi dalam memproduksi produk.",
        "2.	Melakukan produksi sesuai dengan formula dan hanya menggunakan bahan yang tercantum dalam Daftar Bahan Halal. Catatan produksi disimpan setidaknya selama 6 bulan.",
        "3.	Memastikan peralatan produksi dalam keadaan bersih (bebas dari najis) sebelum digunakan."
    )),
    ("Pencucian Fasilitas Produksi", (
        "1.	Peralatan bersih dan kotor dipisahkan",
        "2.	Melakukan proses pencucian dengan bahan pembersih bebas najis",
        "3.	Memastikan proses pencucian menghilangkan kotoran dan najis"
    )),
    ("Penyimpanan & Transportasi", (
        "1.	Menyimpan bahan dan produk di tempat yang bersih dan menjaganya supaya terhindar dari najis.",
        "2.	Memastikan kendaraan yang digunakan untuk mengangkut bahan dan produk halal dalam kondisi baik dan tidak digunakan untuk mengangkut bahan/produk lain yang diragukan kehalalannya."
    )),
    ("Pengembangan Produk Baru", (
        "1.	Setiap ada proses pengembangan produk/menu baru akan didaftarkan untuk disertifikasi ke LPPOM MUI.",
        "2.	Jika sertifikat halal sudah keluar, maka akan dilakukan proses produksi produk/menu baru sesuai dengan formula/resep baku."
    )),
    ("Pengembangan Fasilitas Baru", (
        "1.	Sebelum diajukan untuk mendapatkan sertifikat halal maka pengembangan fasilitas baru dengan terlebih dahulu telah disahkan/telah resmi mendapatkan legalitas untuk semua bentuk perizinan sesuai peraturan daerah/wilayah hukum negara.",
        "2.	Proses pengembangan fasilitas menjamin bahwa pelaksanaan pengembangan fasilitas baru akan didaftarkan sertifikasi halal sebelum dibuka/ digunakan."
    )),
    ("Aturan Karyawan", (
        "1.	Setiap karyawan menjaga kebersihan diri sebelum dan selama bekerja sehingga tidak mengotori produk yang dihasilkan",
        "2.	Setiap karyawan tidak boleh membawa produk tidak halal ke area produksi.",
        "3.	Setiap karyawan tidak boleh membawa/memelihara hewan peliharaan di area produksi.",
        "4.	Setiap karyawan tidak boleh menggunakan peralatan produksi untuk kepentingan lain, misalnya untuk memasak karyawan atau menyimpan produk tidak halal milik karyawan.",
        "5.	Setiap karyawan wajib mencuci tangan sebelum dan sesudah bekerja"
    ))
)

body = doc.add_paragraph(style='List').add_run(f'{prosedur_aktifitas_head}')
body.font.size = Pt(11)
body.font.bold = True


body = doc.add_paragraph(style='List 2').add_run(f'{prosedur_aktifitas_body}')
body.font.size = Pt(11)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Prosedur'
hdr_cells[1].text = 'Langkah Kerja'
for prosedur, langkah_kerja in prosedur_aktifitas_table:
    row_cells = table.add_row().cells
    row_cells[0].text = prosedur
    for langkah in langkah_kerja:
        row_cells[1].text += langkah + "\n"
table.allow_autofit = True
for cell in table.columns[0].cells:
    cell.width = Inches(1)
for cell in table.columns[1].cells:
    cell.width = Inches(5.56)

add_linespace()

kemampuan_telusur_head = "VIII.	 Kemampuan Telusur"
kemampuan_telusur_body = "Perusahaan menjamin produk yang disertifikasi tertelusur berasal dari bahan yang sudah disetujui LPPOM MUI (tercantum dalam Daftar Bahan Halal) dan dibuat di fasilitas produksi yang bebas najis. Kemampuan telusur produk dilakukan melalui catatan pemeriksaan bahan datang dan catatan produksi."

body = doc.add_paragraph(style='List').add_run(f'{kemampuan_telusur_head}')
body.font.size = Pt(11)
body.font.bold = True


body = doc.add_paragraph(style='List 2').add_run(f'{kemampuan_telusur_body}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
body.font.size = Pt(11)

add_linespace()

penanganan_produk_head = "IX.	Penanganan Produk Yang Tidak Memenuhi Kriteria Halal"
penanganan_produk_list = [
    "a.	Produk yang tidak memenuhi kriteria adalah produk yang sudah disertifikasi halal yang dibuat dari bahan yang tidak disetujui (tidak tercantum dalam Daftar Bahan Halal) dan/atau diproduksi di fasilitas yang tidak bebas najis.",
    "b.	Jika ditemukan produk yang tidak memenuhi kriteria halal, maka produk yang dihasilkan tidak akan dijual ke konsumen. Produk tersebut selanjutnya akan dimusnahkan.",
    "c.	Bila produk yang tidak memenuhi kriteria sudah terlanjur dijual, maka produk tersebut akan ditarik dari pasaran dan selanjutnya akan dimusnahkan.",
    "d.	Bukti pemusnahan produk harus disimpan (jika terjadi)."
]

body = doc.add_paragraph(style='List').add_run(f'{penanganan_produk_head}')
body.font.size = Pt(11)
body.font.bold = True

for penanganan_produk in penanganan_produk_list:
    body = doc.add_paragraph(style='List 2').add_run(f'{penanganan_produk}')
    body.font.size = Pt(11)

add_linespace()

audit_internal_head = "X.	Audit Halal Internal"
audit_internal_list = [
    "a.	Perusahaan akan melakukan audit internal dengan cara memeriksa pelaksanaan SJH dan mengisi form seperti pada Lampiran 7. Pertanyaan audit internal ditujukan kepada Tim halal sesuai dengan pembagian tugas Tim halal pada bagian II.",
    "b.	Audit internal dilakukan oleh Tim Halal dan tidak boleh mengaudit tugasnya sendiri.",
    "c.	Audit internal dilakukan minimal 6 bulan sekali.",
    "d.	Jika dalam audit internal ditemukan kelemahan, yaitu ada pertanyaan yang dijawab “tidak”, maka akan segera dilakukan perbaikan agar kelemahan tersebut tidak terulang. Bukti perbaikan kelemahan harus disimpan setidaknya selama dua tahun.",
    "e.	Bukti pelaksanaan audit internal disimpan setidaknya selama dua tahun.",
    "f.	Hasil audit internal dikirimkan ke LPPOM MUI melalui Cerol pada menu Regular Report atau melalui email (sesuai dengan email LPPOM MUI Provinsi)."
]

body = doc.add_paragraph(style='List').add_run(f'{audit_internal_head}')
body.font.size = Pt(11)
body.font.bold = True

for audit_internal in audit_internal_list:
    body = doc.add_paragraph(style='List 2').add_run(f'{audit_internal}')
    body.font.size = Pt(11)

add_linespace()

kaji_ulang_head = "XI.	Kaji Ulang Manajemen"
kaji_ulang_list = [
    "a.	Perusahaan akan melakukan rapat kaji ulang manajemen minimal setahun sekali.",
    "b.	Rapat kaji ulang manajemen dihadiri oleh pimpinan perusahaan dan Tim Manajemen Halal, membahas hasil dari audit internal atau hal lain terkait SJH seperti adanya perubahan tim halal, penambahan produk, pergantian bahan.",
    "c.	Hasil rapat kaji ulang manajemen dituliskan dalam form hasil rapat kaji ulang manajemen seperti pada Lampiran 8."
]

body = doc.add_paragraph(style='List').add_run(f'{kaji_ulang_head}')
body.font.size = Pt(11)
body.font.bold = True

for kaji_ulang in kaji_ulang_list:
    body = doc.add_paragraph(style='List 2').add_run(f'{kaji_ulang}')
    body.font.size = Pt(11)

add_linespace()


# * ============================================================
doc.add_page_break()
# * ============================================================

poster_kebijakan_halal_head = "A.	Poster Kebijakan Halal"
poster_kebijakan_halal_card = {
    "header": "KEBIJAKAN HALAL",
    "company_name": "FAISAL OUTLET",
    "body": '“Kami berkomitmen tinggi untuk menghasilkan produk halal, dengan hanya menggunakan bahan yang telah disetujui oleh LPPOM MUI dan diproduksi dengan menggunakan peralatan yang bebas dari najis. Kami akan mencapainya dengan membentuk tim manajemen halal dan melaksanakan dengan sungguh-sungguh Sistem Jaminan Halal”',
    "date": "Jakarta, 12 Agustus 2021",
    "signature_title": "Pimpinan Perusahaan,",
    "signature_image": "mantap.png",
    "signature_name": "( Faisal Ismail )"
}

body = doc.add_paragraph(style='List').add_run(f'{poster_kebijakan_halal_head}')
body.font.size = Pt(11)
body.font.bold = True

heading_style = doc.styles['Body Text']
head=doc.add_paragraph(style=heading_style).add_run(f'{poster_kebijakan_halal_card["header"]}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
head.font.size = Pt(11)
head.font.bold = True 

heading_style = doc.styles['Body Text']
head=doc.add_paragraph(style=heading_style).add_run(f'{poster_kebijakan_halal_card["company_name"]}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
head.font.size = Pt(11)
head.font.bold = True 


heading_style = doc.styles['Body Text']
head=doc.add_paragraph(style=heading_style).add_run(f'{poster_kebijakan_halal_card["body"]}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
head.font.size = Pt(11)

add_linespace()

heading_style = doc.styles['Body Text']
head=doc.add_paragraph(style=heading_style).add_run(f'{poster_kebijakan_halal_card["date"]}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
head.font.size = Pt(11)

heading_style = doc.styles['Body Text']
head=doc.add_paragraph(style=heading_style).add_run(f'{poster_kebijakan_halal_card["signature_title"]}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
head.font.size = Pt(11)

# signature image & text
ceo_sign = doc.styles['Body Text']
doc.add_picture(poster_kebijakan_halal_card["signature_image"], width=Inches(1))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(style=ceo_sign).add_run(f'{poster_kebijakan_halal_card["signature_name"]}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
ceo_sign.font.size = Pt(11)

add_linespace()

poster_halal_haram_head = "B.	Poster Halal Haram"
poster_halal_haram_card = {
    "header" : "PENGERTIAN HALAL HARAM",
    "list": [
        "	Mengkonsumsi makanan dan minuman yang halal adalah wajib hukumnya bagi orang Islam.",
        "	Pengertian halal haram : (i) Halal adalah Boleh. (ii) Haram adalah sesuatu yang dilarang oleh Allah SWT dengan larangan yang tegas.",
        "	Contoh bahan haram : (i) Babi,termasuk seluruh bagian tubuhnya dan produk turunannya (segar atau olahan), (ii) Khamr (minuman beralkohol), (iii) Hewan sembelihan yang tidak disembelih sesuai syariat Islam, (iv) Darah, (v) Bangkai, (vi) Bagian dari tubuh manusia, (vii) Binatang buas, anjing, amfibi."
    ]
}

add_linespace()

body = doc.add_paragraph(style='List').add_run(f'{poster_halal_haram_head}')
body.font.size = Pt(11)
body.font.bold = True

heading_style = doc.styles['Body Text']
head=doc.add_paragraph(style=heading_style).add_run(f'{poster_halal_haram_card["header"]}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
head.font.size = Pt(11)
head.font.bold = True 

for list_halal_haram in poster_halal_haram_card["list"]:
    body = doc.add_paragraph(style="List 2").add_run(f'{list_halal_haram}')
    body.font.size = Pt(11)


poster_praktek_penerapan_head = "C.	Poster Praktek Penerapan SJH"
poster_praktek_penerapan_card = {
    "header": "PRAKTEK PENERAPAN SJH",
    "list": [
        "	Menjaga fasilitas/peralatan produksi dalam keadaan bersih sebelum dan sesudah digunakan.",
        "	Menjaga kebersihan diri sebelum dan selama bekerja sehingga tidak mengotori produk.",
        "	Tidak boleh membawa produk najis/haram di area produksi.",
        "	Tidak boleh membawa/memelihara hewan peliharaan di area produksi.",
        "	Tidak boleh menggunakan peralatan produksi untuk kepentingan lain.",
        "	Menyimpan bahan dan produk di tempat yang bersih dan menjaga agar terhindar dari najis.",
        "	Memastikan kendaraan yang digunakan untuk mengangkut produk halal dalam kondisi baik dan tidak digunakan untuk mengangkut produk lain yang diragukan kehalalannya."
    ]
}

add_linespace()

body = doc.add_paragraph(style='List').add_run(f'{poster_praktek_penerapan_head}')
body.font.size = Pt(11)
body.font.bold = True

heading_style = doc.styles['Body Text']
head=doc.add_paragraph(style=heading_style).add_run(f'{poster_praktek_penerapan_card["header"]}')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
head.font.size = Pt(11)
head.font.bold = True 

for praktek_penerapan in poster_praktek_penerapan_card["list"]:
    body = doc.add_paragraph(style="List 2").add_run(f'{praktek_penerapan}')
    body.font.size = Pt(11)

add_linespace()


surat_penetapan_tim_head = "SURAT PENETAPAN TIM MANAJEMEN HALAL"
surat_penetapan_tim_text_1 = "Untuk menerapkan Sistem Jaminan Halal dan dalam rangka menjaga konsistensi kehalalan produk, dengan ini ditunjuk Tim Manajemen Halal sebagai berikut :"
surat_penetapan_tim_table = (
    ( "1.", "Faisal Ismail", "Pimpinan", "Ketua"),
    ( "2.", "Eka Dwi Wijayanti", "Produksi", "Anggota"),
    ( "3.", "Muahyatun", "Produksi", "Anggota"),
    ( "4.", "Faisal Ismail", "Penyimpanan & Pengiriman", "Ketua")
)
surat_penetapan_tim_text_2 = "*) Jabatan dapat dirangkap, misalnya 1 orang (nama cukup ditulis satu kali) merangkap sebagai Ketua Tim, Pembelian, Produksi, Gudang dll ."
surat_penetapan_tim_text_3 = "Tim Manajemen Halal telah membaca dan memahami Manual SJH serta akan melaksanakan dengan sungguh-sungguh semua prosedur seperti yang tertulis pada Manual SJH."
surat_penetapan_tim_text_4 = "Demikian surat penetapan ini dibuat untuk dilaksanakan sebagaimana mestinya."

surat_penetapan_tim_signature_date = "Jakarta, 12 Agustus 2021"
surat_penetapan_tim_signature_title = "Pimpinan perusahaan,"
surat_penetapan_tim_signature_image = "mantap.png"
surat_penetapan_tim_signature_name = "( Faisal Ismail )"

# * ============================================================
doc.add_page_break()
# * ============================================================

heading_style = doc.styles['Body Text']
body = doc.add_paragraph(style=heading_style).add_run('Lampiran 2. Surat Penetapan Tim Manajemen Halal')
body.font.size = Pt(11)
body.font.bold = True

# ? ============================================================
# ! Header 
add_linespace()

add_head(surat_penetapan_tim_head)
add_linespace()


heading_style = doc.styles['Body Text']
body = doc.add_paragraph(style=heading_style).add_run(f'{surat_penetapan_tim_text_1}')
body.font.size = Pt(11)


table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No.'
hdr_cells[0].paragraphs[0].runs[0].font.bold = True
hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells[1].text = 'Nama'
hdr_cells[1].paragraphs[0].runs[0].font.bold = True
hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells[2].text = 'Jabatan*)'
hdr_cells[2].paragraphs[0].runs[0].font.bold = True
hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells[3].text = 'Posisi di TIM'
hdr_cells[3].paragraphs[0].runs[0].font.bold = True
hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
for no, nama, jabatan, posisi in surat_penetapan_tim_table:
    row_cells = table.add_row().cells
    row_cells[0].text = no
    row_cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_cells[1].text = nama
    row_cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    row_cells[2].text = jabatan
    row_cells[2].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    row_cells[3].text = posisi
    row_cells[3].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
for cell in table.columns[0].cells:
    cell.width = Inches(0.56)
for cell in table.columns[1].cells:
    cell.width = Inches(2)
for cell in table.columns[2].cells:
    cell.width = Inches(2)
for cell in table.columns[3].cells:
    cell.width = Inches(2)
for row in table.rows:
    row.height = Pt(32)

heading_style = doc.styles['Body Text']
body = doc.add_paragraph(style=heading_style).add_run(f'{surat_penetapan_tim_text_2}')
body.font.size = Pt(11)

heading_style = doc.styles['Body Text']
body = doc.add_paragraph(style=heading_style).add_run(f'{surat_penetapan_tim_text_3}')
body.font.size = Pt(11)

heading_style = doc.styles['Body Text']
body = doc.add_paragraph(style=heading_style).add_run(f'{surat_penetapan_tim_text_4}')
body.font.size = Pt(11)

add_linespace()
add_linespace()

heading_style = doc.styles['Body Text']
body = doc.add_paragraph(style=heading_style).add_run(f'{surat_penetapan_tim_signature_date}')
body.font.size = Pt(11)

heading_style = doc.styles['Body Text']
body = doc.add_paragraph(style=heading_style).add_run(f'{surat_penetapan_tim_signature_title}')
body.font.size = Pt(11)

ceo_sign = doc.styles['Body Text']
doc.add_picture(surat_penetapan_tim_signature_image, width=Inches(1))
doc.add_paragraph(style=ceo_sign).add_run(f'{surat_penetapan_tim_signature_name}')
ceo_sign.font.size = Pt(11)

# * ============================================================
doc.add_page_break()
# * ============================================================

soal_evaluasi_title = "SOAL EVALUASI PELATIHAN INTERNAL SISTEM JAMINAN HALAL"
soal_evaluasi_subtitle = "(soal ini diberikan kepada seluruh peserta pelatihan)"
soal_evaluasi_soal = {

}
soal_evaluasi_jawaban = {
    "nama" : "Dea Jedar",
    "tanggal" : 1657473186975,
    "nilai": '',
    "data" : {
        "1" : "a",
        "2" : "b",
        "3" : "a",
        "4" : "a",
        "5" : "b",
        "6" : "a",
        "7" : "d",
        "8" : "c",
        "9" : True,
        "10" : False
    }
}
pg_styles = doc.styles['List 3']
pg_styles.paragraph_format.line_spacing = 1
pg_styles.paragraph_format.space_after = 0
soal_styles = doc.styles['List 2']
soal_styles.paragraph_format.line_spacing = 1
soal_styles.paragraph_format.space_after = 0
title_soal_style = doc.styles['Body Text']
title_soal_style.paragraph_format.space_after = 0

body = doc.add_paragraph(style=title_soal_style).add_run(f'{soal_evaluasi_title}')
body.font.size = Pt(11)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
body.font.bold = True 

body = doc.add_paragraph(style=title_soal_style).add_run(f'{soal_evaluasi_subtitle}')
body.font.size = Pt(11)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
body.font.bold = True 
body.font.italic = True 

body = doc.add_paragraph(style=title_soal_style).add_run(f'NAMA\t\t: {soal_evaluasi_jawaban["nama"]}')
body.font.size = Pt(11)
body.font.bold = True 

body = doc.add_paragraph(style=title_soal_style).add_run(f'TANGGAL\t: {soal_evaluasi_jawaban["tanggal"]}')
body.font.size = Pt(11)
body.font.bold = True 

body = doc.add_paragraph(style=title_soal_style).add_run(f'NILAI\t\t: {soal_evaluasi_jawaban["nilai"]}')
body.font.size = Pt(11)
body.font.bold = True 


body = doc.add_paragraph(style=soal_styles).add_run('1.	Allah SWT memerintahkan Manusia untuk konsumsi makanan yang…')
body.font.size = Pt(10)

body = doc.add_paragraph(style=pg_styles).add_run('a.	Halal')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['1'] == "a":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('b.	Thoyib ')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['1'] == "b":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('c.	Kotor')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['1'] == "c":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('d.	a dan b')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['1'] == "d":
    body.font.bold = True 



body = doc.add_paragraph(style=soal_styles).add_run('2.	Berikut makanan dan minuman yang halal adalah…')
body.font.size = Pt(10)

body = doc.add_paragraph(style=pg_styles).add_run('a.	Klepon')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['2'] == "a":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('b.	Anjing ')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['2'] == "b":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('c.	Babi')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['2'] == "c":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('d.	Bangkai Ayam')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['2'] == "d":
    body.font.bold = True 



body = doc.add_paragraph(style=soal_styles).add_run('3.	Daging Babi dan turunannya merupakan najis…')
body.font.size = Pt(10)

body = doc.add_paragraph(style=pg_styles).add_run('a.	Ringan (mukhaffafah)')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['3'] == "a":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('b.	Berat (mughallazhah)')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['3'] == "b":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('c.	Sedang (mutawassithah)')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['3'] == "c":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('d.	Tidak Najis')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['3'] == "d":
    body.font.bold = True 



body = doc.add_paragraph(style=soal_styles).add_run('4.	Cara Mengghilangkan Najis Sedang (mutawassithah) yaitu…')
body.font.size = Pt(10)

body = doc.add_paragraph(style=pg_styles).add_run('a.	Dengan mengucurinya dengan air atau mencucinya di dalam air yang banyak (direndam) hingga hilang rasa, bau dan warna dari bahan najisnya.')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['4'] == "a":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('b.	Di Diamkan Saja')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['4'] == "b":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('c.	Di Bakar')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['4'] == "c":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('d.	Dicuci tujuh kali dengan air dan salah satunya dengan tanah atau bahan lain yang mempunyai kemampuan menghilangkan rasa, bau dan warna.')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['4'] == "d":
    body.font.bold = True 



body = doc.add_paragraph(style=soal_styles).add_run('5.	Cara Mengghilangkan Najis Berat (mughallazhah) yaitu…')
body.font.size = Pt(10)

body = doc.add_paragraph(style=pg_styles).add_run('a.	Dengan mengucurinya dengan air atau mencucinya di dalam air yang banyak (direndam) hingga hilang rasa, bau dan warna dari bahan najisnya.')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['5'] == "a":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('b.	Di Diamkan Saja')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['5'] == "b":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('c.	Di Bakar')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['5'] == "c":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('d.	Dicuci tujuh kali dengan air dan salah satunya dengan tanah atau bahan lain yang mempunyai kemampuan menghilangkan rasa, bau dan warna.')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['5'] == "d":
    body.font.bold = True 



body = doc.add_paragraph(style=soal_styles).add_run('6.	Cara menjaga Konsistensi dalam memproduksi Produk dan bahan yang halal perusahaan harus menerapkan…')
body.font.size = Pt(10)

body = doc.add_paragraph(style=pg_styles).add_run('a.	Sistem Keamanan Pangan')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['6'] == "a":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('b.	Sistem Keselamatan Kerja')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['6'] == "b":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('c.	Sistem Jaminan Halal')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['6'] == "c":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('d.	Sistem Informasi')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['6'] == "d":
    body.font.bold = True 



body = doc.add_paragraph(style=soal_styles).add_run('7.	Kriteria Sistem Jaminan Halal Terdiri dari …. Kriteria')
body.font.size = Pt(10)

body = doc.add_paragraph(style=pg_styles).add_run('a.	11')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['7'] == "a":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('b.	20')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['7'] == "b":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('c.	12')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['7'] == "c":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('d.	14')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['7'] == "d":
    body.font.bold = True 



body = doc.add_paragraph(style=soal_styles).add_run('8.	Aktifitas manakah dibawah ini yang merupakan aktifitas kritis dalam Sistem Jaminan Halal')
body.font.size = Pt(10)

body = doc.add_paragraph(style=pg_styles).add_run('a.	Seleksi Bahan Baru')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['8'] == "a":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('b.	Pembelian')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['8'] == "b":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('c.	Formulasi Produk baru')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['8'] == "c":
    body.font.bold = True 

body = doc.add_paragraph(style=pg_styles).add_run('d.	Semua Benar')
body.font.size = Pt(10)
if soal_evaluasi_jawaban['data']['8'] == "d":
    body.font.bold = True 

body = doc.add_paragraph(style=soal_styles)
soal = body.add_run('9.	Setiap ada Bahan Baru perusahaan tidak wajib melaporkan kepada LPPOM MUI ')
soal.font.size = Pt(10)
if soal_evaluasi_jawaban["data"]['9'] == True:
    non_slash = body.add_run("(Benar/")
    non_slash.font.bold = True 
    slash = body.add_run("Salah)")
    slash.font.strike = True
    slash.font.bold = True 
    slash.font.size = Pt(10)
    non_slash.font.size = Pt(10)
else:
    slash = body.add_run("(Benar")
    slash.font.strike = True
    slash.font.bold = True 
    slash.font.size = Pt(10)
    non_slash = body.add_run("/Salah)")
    non_slash.font.bold = True 
    non_slash.font.size = Pt(10)


body = doc.add_paragraph(style=soal_styles)
soal = body.add_run('10.	Audit Internal dilakukan 6 Bulan Sekali dan dilaporkan kepada LPPOM MUI')
soal.font.size = Pt(10)
if soal_evaluasi_jawaban["data"]['10'] == True:
    non_slash = body.add_run("(Benar/")
    non_slash.font.bold = True 
    slash = body.add_run("Salah)")
    slash.font.strike = True
    slash.font.bold = True 
    slash.font.size = Pt(10)
    non_slash.font.size = Pt(10)
else:
    slash = body.add_run("(Benar")
    slash.font.strike = True
    slash.font.bold = True 
    slash.font.size = Pt(10)
    non_slash = body.add_run("/Salah)")
    non_slash.font.bold = True 
    non_slash.font.size = Pt(10)
doc.save("Uji.docx")